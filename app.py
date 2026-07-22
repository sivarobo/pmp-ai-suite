import streamlit as st
import pandas as pd
from google import genai
from google.genai import types
from PIL import Image
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import io
import re
import time
import datetime
import random
import json
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np
import psycopg2
import psycopg2.extras
from streamlit_oauth import OAuth2Component
import requests

# ==========================================
# st.set_page_config - MUST BE FIRST
# ==========================================
st.set_page_config(page_title="PMP Edu AI", page_icon="🎓", layout="wide", initial_sidebar_state="collapsed")

# ==========================================
# CSS Styling
# ==========================================
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Sora:wght@600;700;800&family=Inter:wght@400;500;600;700&family=Noto+Sans+Tamil:wght@400;500;600;700&display=swap');

    :root{
        --navy:#0a1f44; --navy2:#12305f; --navy3:#1a3a6b;
        --gold:#c9a227; --gold-lt:#e6c866; --gold-soft:#f3e7bd;
        --bg:#f4f6fb; --card:#ffffff; --ink:#0f1a30; --slate:#5a6782; --line:#e6e9f2;
    }

    /* ===== Global ===== */
    .stApp { background: var(--bg); }
    .block-container { padding-top: 3rem !important; padding-bottom: 1rem !important; max-width: 1150px !important; }
    /* tighten vertical gaps between widgets */
    [data-testid="stVerticalBlock"] { gap: 0.35rem !important; }
    div[data-testid="stHorizontalBlock"] { gap: 0.6rem !important; }
    html, body, [class*="css"] { font-family: 'Inter','Noto Sans Tamil',sans-serif; }
    [data-testid="stMarkdownContainer"] p { font-size: 14.5px; font-weight: 500; color: var(--ink); }

    /* compact headings */
    h1 { font-size: 22px !important; }
    h2 { font-size: 18px !important; }
    h3 { font-size: 15px !important; }
    h4 { font-size: 14px !important; }

    /* compact inputs */
    .stSelectbox div[data-baseweb="select"] > div,
    .stTextInput input, .stNumberInput input {
        border-radius: 9px !important; border: 1.5px solid var(--line) !important;
        background: #fbfcfe !important; font-size: 14px !important;
        padding-top: 6px !important; padding-bottom: 6px !important; min-height: 38px !important;
    }
    .stSelectbox label, .stTextInput label, .stNumberInput label, .stRadio label, .stMultiSelect label, .stSlider label {
        font-size: 12.5px !important; font-weight: 600 !important; color: var(--slate) !important;
        margin-bottom: 2px !important; padding-bottom: 0 !important;
    }
    /* compact number input +/- buttons */
    .stNumberInput button { min-height: 38px !important; height: 38px !important; padding: 0 10px !important; }

    /* stronger visible borders on all inputs */
    .stTextInput input, .stNumberInput input,
    .stSelectbox div[data-baseweb="select"] > div,
    .stMultiSelect div[data-baseweb="select"] > div {
        border: 1.5px solid #c9cfdd !important; box-shadow: 0 1px 2px rgba(10,31,68,.04) !important;
    }
    .stTextInput input:focus, .stNumberInput input:focus { border-color: var(--gold) !important; }

    /* Section card wrapper */
    .pmp-section {
        background: #ffffff; border: 1.5px solid #d9dfea; border-radius: 16px;
        padding: 16px 18px; margin: 10px 0; box-shadow: 0 2px 8px rgba(10,31,68,.05);
    }

    /* Headings */
    h1, h2, h3 { font-family: 'Sora','Noto Sans Tamil',sans-serif !important; color: var(--navy) !important; font-weight: 800 !important; }

    /* ===== Sidebar ===== */
    [data-testid="stSidebar"] {
        background: linear-gradient(180deg, var(--navy), #081934) !important;
    }
    [data-testid="stSidebar"] * { color: #c7d1e6 !important; }
    [data-testid="stSidebar"] h1, [data-testid="stSidebar"] h2, [data-testid="stSidebar"] h3 { color: #ffffff !important; }
    [data-testid="stSidebar"] .stButton>button {
        background: rgba(255,255,255,.07) !important;
        color: #e8ecf4 !important;
        border: 1px solid rgba(201,162,39,.25) !important;
    }
    [data-testid="stSidebar"] .stButton>button:hover {
        background: linear-gradient(135deg, var(--gold), var(--gold-lt)) !important;
        color: var(--navy) !important;
        border-color: var(--gold) !important;
    }

    /* ===== Buttons (Gold primary) ===== */
    .stButton>button, .stDownloadButton>button, .stFormSubmitButton>button {
        border-radius: 10px !important; font-weight: 700 !important; font-size: 14px !important;
        height: auto !important; padding: 9px 18px !important; transition: .2s !important;
        border: 1.5px solid var(--line) !important; background: #ffffff !important; color: var(--navy) !important;
    }
    .stButton>button[kind="primary"], .stDownloadButton>button, .stFormSubmitButton>button {
        background: linear-gradient(135deg, var(--gold), var(--gold-lt)) !important;
        color: var(--navy) !important; border: none !important;
        box-shadow: 0 8px 22px rgba(201,162,39,.35) !important;
    }
    .stButton>button:hover, .stDownloadButton>button:hover, .stFormSubmitButton>button:hover {
        transform: translateY(-2px) !important;
    }

    /* ===== Tabs ===== */
    .stTabs [data-baseweb="tab-list"] { gap: 8px; border-bottom: none; }
    .stTabs [data-baseweb="tab"] {
        background: #ffffff; border: 1px solid var(--line); border-radius: 11px;
        padding: 9px 18px; font-weight: 600; color: var(--slate);
    }
    .stTabs [aria-selected="true"] {
        background: linear-gradient(135deg, var(--navy), var(--navy3)) !important;
        border-color: var(--navy) !important;
    }
    .stTabs [aria-selected="true"] * {
        color: #ffffff !important;
    }

    /* ===== Expander as card ===== */
    .streamlit-expanderHeader, [data-testid="stExpander"] {
        background: var(--card) !important; border-radius: 14px !important; border: 1px solid var(--line) !important;
    }

    /* ===== Metrics / cards ===== */
    [data-testid="stMetric"] {
        background: var(--card); border: 1px solid var(--line); border-radius: 16px;
        padding: 18px 20px; box-shadow: 0 2px 6px rgba(10,31,68,.05);
    }
    [data-testid="stMetricValue"] { font-family:'Sora',sans-serif; color: var(--navy); font-weight: 800; }

    /* Radio horizontal pills */
    .stRadio [role="radiogroup"] { gap: 10px; }

    /* Alerts softer */
    .stAlert { border-radius: 12px; }

    /* ===== Custom dashboard header ===== */
    .pmp-header {
        background: linear-gradient(135deg, var(--navy), var(--navy3));
        border-radius: 18px; padding: 26px 30px; margin-bottom: 22px;
        display: flex; justify-content: space-between; align-items: center; flex-wrap: wrap; gap: 14px;
        border-bottom: 3px solid var(--gold);
    }
    .pmp-header h1 { color: #fff !important; font-size: 24px !important; margin: 0 !important; }
    .pmp-header h1 .accent { color: var(--gold-lt) !important; }
    .pmp-header p { color: #b9c4dd !important; font-size: 14px; margin: 4px 0 0 0; }
    .pmp-badge { background: var(--gold-soft); border: 1.5px solid var(--gold); color: var(--navy);
                 padding: 8px 16px; border-radius: 10px; font-weight: 700; font-size: 13px; }

    /* Google Login button (kept) */
    .google-btn-link {
        display:inline-flex; align-items:center; justify-content:center; gap:12px;
        background:#fff; color:#3c4043; border:1.5px solid #dadce0; border-radius:12px;
        padding:14px 32px; font-size:16px; font-weight:600; text-decoration:none !important;
        width:100%; box-shadow:0 2px 8px rgba(0,0,0,0.10); transition:.2s;
    }
    .google-btn-link:hover { box-shadow:0 6px 18px rgba(201,162,39,0.25); background:#fffdf5; }
</style>
""", unsafe_allow_html=True)

# ==========================================
# DB Connection (Neon PostgreSQL)
# ==========================================
def get_db():
    try:
        conn = psycopg2.connect(
            st.secrets["NEON_DATABASE_URL"],
            cursor_factory=psycopg2.extras.RealDictCursor
        )
        return conn
    except Exception as e:
        st.error(f"DB Connection Error: {e}")
        return None

# ==========================================
# Google OAuth — streamlit-oauth library
# ==========================================
GOOGLE_CLIENT_ID     = st.secrets["google"]["client_id"]
GOOGLE_CLIENT_SECRET = st.secrets["google"]["client_secret"]
GOOGLE_REDIRECT_URI  = st.secrets["google"]["redirect_uri"]

oauth2 = OAuth2Component(
    client_id=GOOGLE_CLIENT_ID,
    client_secret=GOOGLE_CLIENT_SECRET,
    authorize_endpoint="https://accounts.google.com/o/oauth2/v2/auth",
    token_endpoint="https://oauth2.googleapis.com/token",
    refresh_token_endpoint="https://oauth2.googleapis.com/token",
    revoke_token_endpoint="https://oauth2.googleapis.com/revoke",
)

def _google_userinfo(access_token):
    r = requests.get(
        "https://www.googleapis.com/oauth2/v2/userinfo",
        headers={"Authorization": f"Bearer {access_token}"},
        timeout=10
    )
    return r.json() if r.status_code == 200 else None

# ==========================================
# User DB Functions (Google-based)
# ==========================================
def upsert_google_user(google_info):
    """
    Google login → DB-ல் user save/update.
    streamlit-google-auth provides: name, email, picture
    """
    try:
        conn = get_db()
        if not conn:
            return None
        cur = conn.cursor()

        # Table create (first run only - safe to run every time)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS users (
                id          SERIAL PRIMARY KEY,
                google_id   TEXT UNIQUE,
                email       TEXT UNIQUE,
                name        TEXT,
                picture     TEXT,
                plan        TEXT DEFAULT 'free',
                daily_count INTEGER DEFAULT 0,
                last_used   DATE DEFAULT CURRENT_DATE,
                created_at  TIMESTAMP DEFAULT NOW(),
                last_login  TIMESTAMP DEFAULT NOW(),
                school_name  TEXT,
                teacher_name TEXT,
                mobile       TEXT
            )
        """)
        # Safe migration for existing DBs (columns added if missing)
        for col in ["google_id TEXT", "picture TEXT", "plan TEXT DEFAULT 'free'",
                    "created_at TIMESTAMP DEFAULT NOW()", "last_login TIMESTAMP DEFAULT NOW()",
                    "school_name TEXT", "teacher_name TEXT", "mobile TEXT"]:
            try:
                cur.execute(f"ALTER TABLE users ADD COLUMN IF NOT EXISTS {col}")
            except Exception:
                pass
        conn.commit()

        # Upsert by email (google_id optional)
        cur.execute("""
            INSERT INTO users (email, name, picture, last_login)
            VALUES (%s, %s, %s, NOW())
            ON CONFLICT (email) DO UPDATE
              SET name       = EXCLUDED.name,
                  picture    = EXCLUDED.picture,
                  last_login = NOW()
            RETURNING id, email, name, picture, plan, created_at, school_name, teacher_name, mobile
        """, (
            google_info["email"],
            google_info["name"],
            google_info.get("picture", ""),
        ))

        row = cur.fetchone()
        conn.commit()
        conn.close()
        return dict(row) if row else None

    except Exception as e:
        st.error(f"DB Error: {e}")
        return None

def fetch_user_by_email(email):
    try:
        conn = get_db()
        if not conn:
            return None
        cur = conn.cursor()
        cur.execute(
            "SELECT id, email, name, picture, plan, created_at, school_name, teacher_name, mobile "
            "FROM users WHERE email=%s", (email,)
        )
        row = cur.fetchone()
        conn.close()
        return dict(row) if row else None
    except Exception as e:
        st.warning(f"⚠️ User fetch பிழை: {e}")
        return None

def update_user_profile(user_id, school_name, teacher_name, mobile, email=None):
    try:
        conn = get_db()
        if not conn:
            return False
        cur = conn.cursor()
        # Prefer email as the key — robust whether id is serial/uuid or missing.
        if email:
            cur.execute(
                "UPDATE users SET school_name=%s, teacher_name=%s, mobile=%s WHERE email=%s",
                (school_name.strip(), teacher_name.strip(), mobile.strip(), email)
            )
        else:
            cur.execute(
                "UPDATE users SET school_name=%s, teacher_name=%s, mobile=%s WHERE id=%s",
                (school_name.strip(), teacher_name.strip(), mobile.strip(), user_id)
            )
        conn.commit()
        conn.close()
        return True
    except Exception as e:
        st.error(f"Profile Save Error: {e}")
        return False

def get_today_usage(user_id):
    try:
        conn = get_db()
        if not conn:
            return 0
        cur = conn.cursor()
        cur.execute(
            "SELECT question_count FROM daily_usage WHERE user_id = %s AND usage_date = CURRENT_DATE",
            (str(user_id),)
        )
        row = cur.fetchone()
        conn.close()
        return row["question_count"] if row else 0
    except:
        return 0

def increment_usage(user_id):
    try:
        conn = get_db()
        if not conn:
            return False
        cur = conn.cursor()
        cur.execute(
            """INSERT INTO daily_usage (user_id, usage_date, question_count)
               VALUES (%s, CURRENT_DATE, 1)
               ON CONFLICT (user_id, usage_date)
               DO UPDATE SET question_count = daily_usage.question_count + 1""",
            (str(user_id),)
        )
        conn.commit()
        conn.close()
        return True
    except:
        return False

# ==========================================
# QUESTION BANK (புத்தக கேள்வி வங்கி) — DB Functions
# ==========================================
def ensure_question_bank_table():
    try:
        conn = get_db()
        if not conn:
            return
        cur = conn.cursor()
        cur.execute("""
            CREATE TABLE IF NOT EXISTS question_bank (
                id            SERIAL PRIMARY KEY,
                subject       TEXT,
                lesson        TEXT,
                mark_type     TEXT,
                qtype         TEXT,
                question_text TEXT,
                answer_text   TEXT,
                reference     TEXT,
                created_at    TIMESTAMP DEFAULT NOW()
            )
        """)
        try:
            cur.execute("ALTER TABLE question_bank ADD COLUMN IF NOT EXISTS reference TEXT")
        except Exception:
            pass
        conn.commit()
        conn.close()
    except Exception as e:
        st.error(f"Question Bank Table Error: {e}")

def fetch_bank_questions(subject, lesson, mark_type):
    try:
        conn = get_db()
        if not conn:
            return []
        cur = conn.cursor()
        cur.execute(
            "SELECT id, qtype, question_text, answer_text, reference FROM question_bank "
            "WHERE subject=%s AND lesson=%s AND mark_type=%s ORDER BY id",
            (subject, lesson, mark_type)
        )
        rows = cur.fetchall()
        conn.close()
        return [dict(r) for r in rows]
    except Exception as e:
        st.error(f"Bank Fetch Error: {e}")
        return []

MARK_TYPE_LABELS = {
    "1M":   "1-மார்க் (MCQ / குறுவினா)",
    "2M":   "2-மார்க் வினா",
    "5M":   "5-மார்க் வினா",
    "LONG": "நெடுவினா (Long Answer)",
}

def save_bank_questions(subject, lesson, mark_type, items):
    try:
        conn = get_db()
        if not conn:
            return
        cur = conn.cursor()
        for it in items:
            cur.execute(
                "INSERT INTO question_bank (subject, lesson, mark_type, qtype, question_text, answer_text) "
                "VALUES (%s, %s, %s, %s, %s, %s)",
                (subject, lesson, mark_type, it.get("qtype", "பயிற்சி"), it.get("question", "").strip(), it.get("answer", "").strip())
            )
        conn.commit()
        conn.close()
    except Exception as e:
        st.error(f"Bank Save Error: {e}")

def generate_bank_questions_ai(subject, lesson, mark_type, count=8):
    """Gemini மூலம் புத்தக பாணி (Book Back / Exercise / Example) கேள்விகளை உருவாக்குதல்"""
    mark_desc = MARK_TYPE_LABELS.get(mark_type, mark_type)
    prompt = f"""நீங்கள் Tamil Nadu Class 10 பாடநூல் நிபுணர். பாடம்: "{lesson}" | பாடப்பிரிவு: {subject}.

இந்த பாடத்திலிருந்து {count} கேள்விகளை உருவாக்கவும் — மதிப்பெண் வகை: {mark_desc}.
இவை TN Samacheer Kalvi பாடப்புத்தகத்தின் "பின்புற வினாக்கள்" (Book Back Exercise), "பயிற்சி கணக்குகள்" (Practice Problems),
மற்றும் "எடுத்துக்காட்டு கணக்குகள்" (Solved Examples) பாணியை நெருக்கமாக ஒத்திருக்க வேண்டும்.
qtype-ஐ மூன்றிலும் கலந்து (mix) கொடுக்கவும்.

பதில் STRICTLY இந்த JSON அணி (array) வடிவில் மட்டும் இருக்க வேண்டும், வேறு எந்த உரையும் (preamble, code fences) கூடாது:
[
  {{"qtype": "பின்புற வினா", "question": "தமிழில் முழு கேள்வி", "answer": "தமிழில் சுருக்கமான விடை/தீர்வு"}},
  ...
]
கணிதம் சின்னங்களுக்கு LaTeX பயன்படுத்தாதீர்கள் — plain Unicode (×, ÷, √, ², π, ∠ போன்றவை) மட்டும் பயன்படுத்தவும்.
"""
    try:
        response = gemini_generate(prompt)
        raw = response.text.strip()
        raw = re.sub(r'^```json\s*|\s*```$', '', raw.strip(), flags=re.MULTILINE).strip()
        raw = re.sub(r'^```|```$', '', raw.strip()).strip()
        items = json.loads(raw)
        if isinstance(items, list):
            return items
        return []
    except Exception as e:
        emsg = str(e)
        if "429" in emsg or "RESOURCE_EXHAUSTED" in emsg:
            st.warning(f"⚠️ '{lesson}' — Gemini free-tier தினசரி வரம்பு (20 requests) முடிந்துவிட்டது. கொஞ்ச நேரம் கழித்து முயற்சிக்கவும், அல்லது ஏற்கனவே Import செய்த கேள்வி வங்கியை பயன்படுத்தவும்.")
        else:
            st.warning(f"⚠️ '{lesson}' ({mark_desc}) கேள்வி வங்கி உருவாக்கத்தில் பிழை: {e}")
        return []

def get_or_build_bank(subject, lesson, mark_type, min_count=8):
    """கேள்வி வங்கியில் இருந்து மட்டும் fetch பண்ணும். AI-ஐ ஒருபோதும் கூப்பிடாது.
       (Imported புத்தக கேள்விகளை மட்டும் பயன்படுத்துகிறோம் — quota பிரச்சனை வராது.)"""
    existing = fetch_bank_questions(subject, lesson, mark_type)
    if existing:
        return existing
    # Subject/lesson பெயர் exact match ஆகவில்லை எனில், case-insensitive-ஆ மீண்டும் முயற்சி
    try:
        conn = get_db()
        if conn:
            cur = conn.cursor()
            cur.execute(
                "SELECT id, qtype, question_text, answer_text, reference FROM question_bank "
                "WHERE LOWER(TRIM(lesson))=LOWER(TRIM(%s)) AND mark_type=%s ORDER BY id",
                (lesson, mark_type)
            )
            rows = cur.fetchall()
            conn.close()
            if rows:
                return [dict(r) for r in rows]
    except Exception:
        pass
    return []

def assemble_paper_from_bank(parts_cfg):
    """
    parts_cfg: list of dicts —
      {"label": "பகுதி I", "mark": 1, "given": N, "answer": N, "note": "", "items": [ {question_text, answer_text}, ... ]}
    தேர்ந்தெடுக்கப்பட்ட கேள்விகளை மட்டும் வைத்து AI-response போன்ற text உருவாக்குகிறது
    (create_professional_docx() அதே pipeline-ஐ மறுபயன்பாடு செய்ய).
    """
    q_lines = []
    a_lines = []
    counter = 1
    for part in parts_cfg:
        items = part["items"]
        if not items:
            continue
        total_marks = part["given"] * part["mark"]
        header = f'{part["label"]} ( {part["given"]} x {part["mark"]} = {total_marks} )'
        if part.get("note"):
            header += f'  [{part["note"]}]'
        q_lines.append(header)
        a_lines.append(part["label"])
        for it in items:
            q_lines.append(f'{counter}. {it["question_text"]}')
            a_lines.append(f'{counter}. {it.get("answer_text", "")}')
            counter += 1
        q_lines.append("")
        a_lines.append("")
    return "\n".join(q_lines) + "\n=== ANSWER KEY ===\n" + "\n".join(a_lines)

def update_bank_question(qid, question_text, answer_text, qtype):
    try:
        conn = get_db()
        if not conn:
            return False
        cur = conn.cursor()
        cur.execute(
            "UPDATE question_bank SET question_text=%s, answer_text=%s, qtype=%s WHERE id=%s",
            (question_text, answer_text, qtype, qid)
        )
        conn.commit()
        conn.close()
        return True
    except Exception as e:
        st.error(f"Update Error: {e}")
        return False

def delete_bank_question(qid):
    try:
        conn = get_db()
        if not conn:
            return False
        cur = conn.cursor()
        cur.execute("DELETE FROM question_bank WHERE id=%s", (qid,))
        conn.commit()
        conn.close()
        return True
    except Exception as e:
        st.error(f"Delete Error: {e}")
        return False

def list_bank_subjects():
    try:
        conn = get_db()
        if not conn:
            return []
        cur = conn.cursor()
        cur.execute("SELECT DISTINCT subject FROM question_bank ORDER BY subject")
        rows = cur.fetchall()
        conn.close()
        return [r["subject"] for r in rows]
    except Exception:
        return []

def list_bank_lessons(subject):
    try:
        conn = get_db()
        if not conn:
            return []
        cur = conn.cursor()
        cur.execute("SELECT DISTINCT lesson FROM question_bank WHERE subject=%s ORDER BY lesson", (subject,))
        rows = cur.fetchall()
        conn.close()
        return [r["lesson"] for r in rows]
    except Exception:
        return []

def fetch_bank_filtered(subject=None, lesson=None, mark_type=None):
    try:
        conn = get_db()
        if not conn:
            return []
        cur = conn.cursor()
        query = "SELECT id, subject, lesson, mark_type, qtype, question_text, answer_text FROM question_bank WHERE 1=1"
        params = []
        if subject:
            query += " AND subject=%s"; params.append(subject)
        if lesson:
            query += " AND lesson=%s"; params.append(lesson)
        if mark_type:
            query += " AND mark_type=%s"; params.append(mark_type)
        query += " ORDER BY subject, lesson, mark_type, id"
        cur.execute(query, tuple(params))
        rows = cur.fetchall()
        conn.close()
        return [dict(r) for r in rows]
    except Exception as e:
        st.error(f"Fetch Error: {e}")
        return []

def bulk_import_bank_from_df(df_import):
    """
    Excel template columns: Subject, Lesson, Mark_Type, QType, Question, Answer
    ஏற்கனவே DB-ல் இல்லாத (Subject+Lesson+Mark_Type+Question) rows-ஐ மட்டும் insert பண்ணும்.
    """
    required_cols = {"Subject", "Lesson", "Mark_Type", "QType", "Question", "Answer"}
    if not required_cols.issubset(set(df_import.columns)):
        return 0, 0, f"தேவையான Columns இல்லை. Template-ல் இருக்கும் columns: {', '.join(sorted(required_cols))}"
    inserted, skipped = 0, 0
    try:
        conn = get_db()
        if not conn:
            return 0, 0, "DB Connection தோல்வி"
        cur = conn.cursor()
        for _, row in df_import.iterrows():
            subject = str(row["Subject"]).strip()
            lesson  = str(row["Lesson"]).strip()
            mtype   = str(row["Mark_Type"]).strip().upper()
            qtype   = str(row["QType"]).strip()
            qtext   = str(row["Question"]).strip()
            atext   = str(row["Answer"]).strip() if not pd.isna(row["Answer"]) else ""
            ref     = str(row["Reference"]).strip() if ("Reference" in df_import.columns and not pd.isna(row.get("Reference"))) else ""
            if not subject or not lesson or not qtext or subject == "nan" or qtext == "nan":
                skipped += 1
                continue
            cur.execute(
                "SELECT id FROM question_bank WHERE subject=%s AND lesson=%s AND mark_type=%s AND question_text=%s",
                (subject, lesson, mtype, qtext)
            )
            if cur.fetchone():
                skipped += 1
                continue
            cur.execute(
                "INSERT INTO question_bank (subject, lesson, mark_type, qtype, question_text, answer_text, reference) "
                "VALUES (%s, %s, %s, %s, %s, %s, %s)",
                (subject, lesson, mtype, qtype, qtext, atext, ref)
            )
            inserted += 1
        conn.commit()
        conn.close()
        return inserted, skipped, None
    except Exception as e:
        return inserted, skipped, str(e)

def fetch_bank_missing_answers(subject=None, lesson=None, mark_type=None, limit=500):
    try:
        conn = get_db()
        if not conn:
            return []
        cur = conn.cursor()
        query = ("SELECT id, subject, lesson, mark_type, qtype, question_text FROM question_bank "
                  "WHERE (answer_text IS NULL OR TRIM(answer_text) = '')")
        params = []
        if subject:
            query += " AND subject=%s"; params.append(subject)
        if lesson:
            query += " AND lesson=%s"; params.append(lesson)
        if mark_type:
            query += " AND mark_type=%s"; params.append(mark_type)
        query += " ORDER BY id LIMIT %s"
        params.append(limit)
        cur.execute(query, tuple(params))
        rows = cur.fetchall()
        conn.close()
        return [dict(r) for r in rows]
    except Exception as e:
        st.error(f"Fetch Error: {e}")
        return []

def generate_answers_batch_ai(items, batch_size=10):
    """
    items: list of dicts with id, subject, lesson, question_text
    Gemini-ஐ batch-ஆ கூப்பிட்டு, ஒவ்வொரு கேள்விக்கும் சுருக்கமான solution வாங்குது.
    Returns dict[id] = answer_text
    """
    results = {}
    for i in range(0, len(items), batch_size):
        chunk = items[i:i + batch_size]
        q_list = "\n".join([f'{idx+1}. [{it["lesson"]}] {it["question_text"]}' for idx, it in enumerate(chunk)])
        prompt = f"""நீங்கள் TN Samacheer Kalvi Class 10 பாட நிபுணர். கீழே {len(chunk)} கணக்கு/கேள்விகள் கொடுக்கப்பட்டுள்ளன.
ஒவ்வொன்றுக்கும் சுருக்கமான, தெளிவான Solution/Answer (தமிழ் அல்லது English medium-க்கு ஏற்ப) கொடுக்கவும்.

{q_list}

பதில் STRICTLY இந்த JSON array வடிவில் மட்டும் இருக்க வேண்டும் (வேறு உரை/code fence கூடாது), அதே வரிசையில்:
[
  {{"n": 1, "answer": "..."}},
  {{"n": 2, "answer": "..."}}
]
கணிதச் சின்னங்களுக்கு Plain Unicode (×, ÷, √, ², π, ∠, ∈, ∪, ∩, ≤, ≥, ⇒) பயன்படுத்தவும், LaTeX வேண்டாம்.
"""
        try:
            response = gemini_generate(prompt)
            raw = response.text.strip()
            raw = re.sub(r'^```json\s*|\s*```$', '', raw, flags=re.MULTILINE).strip()
            raw = re.sub(r'^```|```$', '', raw.strip()).strip()
            parsed = json.loads(raw)
            for item in parsed:
                n = item.get("n")
                if n and 1 <= n <= len(chunk):
                    results[chunk[n - 1]["id"]] = item.get("answer", "").strip()
        except Exception as e:
            emsg = str(e)
            if "429" in emsg or "RESOURCE_EXHAUSTED" in emsg:
                st.warning("⚠️ Gemini free-tier தினசரி வரம்பு முடிந்தது. மீதி batches-ஐ நாளை/கொஞ்ச நேரம் கழித்து முயற்சிக்கவும்.")
                break
            st.warning(f"⚠️ Batch {i//batch_size + 1} பிழை: {e}")
    return results

FREE_DAILY_LIMIT = 2

# ==========================================
# GOOGLE OAUTH CALLBACK HANDLER
# ==========================================
# ==========================================
# ACCESS GATE — streamlit-oauth
# ==========================================
if not st.session_state.get("logged_in_user"):
    # ===== INTRO BANNER SPLASH (before login, once per session) =====
    if not st.session_state.get("intro_shown"):
        import os as _os_i, base64 as _b64_i
        _bn_i = None
        for _bf in ("banner.png", "banner.jpg", "banner.jpeg", "banner.webp"):
            if _os_i.path.exists(_bf):
                _bn_i = _bf
                break
        _bhtml = ""
        if _bn_i:
            try:
                with open(_bn_i, "rb") as _f:
                    _b64 = _b64_i.b64encode(_f.read()).decode()
                _ext = _bn_i.rsplit(".", 1)[-1]
                _bhtml = (f"<img src='data:image/{_ext};base64,{_b64}' "
                          f"style='max-width:88%;max-height:76vh;border-radius:20px;"
                          f"box-shadow:0 26px 70px rgba(0,0,0,.45);'/>")
            except Exception:
                _bhtml = ""
        if not _bhtml:
            _bhtml = ("<div style='font-family:Sora,sans-serif;font-size:46px;font-weight:800;color:#fff;'>"
                      "PMP <span style='color:#e6c866;'>Edu</span> AI</div>"
                      "<div style='color:#c7d1e6;font-size:16px;margin-top:10px;'>Smart Questions · Better Learning</div>")

        st.markdown(f"""
        <style>
            [data-testid="stHeader"] {{ display:none; }}
            .block-container {{ padding:0 !important; }}
            #intro {{
                position:fixed; inset:0; z-index:99999;
                background:linear-gradient(135deg,#0a1f44,#12305f);
                display:flex; flex-direction:column; align-items:center; justify-content:center;
                animation: introFade 4s ease-in-out forwards;
            }}
            #intro .inner {{ text-align:center; animation: introZoom 1.1s cubic-bezier(.2,.8,.2,1); }}
            #intro .bar {{ margin-top:22px; width:200px; height:5px; background:rgba(255,255,255,.15);
                           border-radius:5px; overflow:hidden; }}
            #intro .bar::after {{ content:""; display:block; height:100%; width:0;
                           background:linear-gradient(90deg,#3b82f6,#a855f7,#f97316);
                           border-radius:5px; animation: introBar 4s linear forwards; }}
            @keyframes introFade {{ 0%{{opacity:0;}} 8%{{opacity:1;}} 82%{{opacity:1;}} 100%{{opacity:0; visibility:hidden;}} }}
            @keyframes introZoom {{ 0%{{transform:scale(.88); opacity:0;}} 100%{{transform:scale(1); opacity:1;}} }}
            @keyframes introBar {{ 0%{{width:0;}} 100%{{width:100%;}} }}
        </style>
        <div id="intro"><div class="inner">{_bhtml}<div class="bar"></div></div></div>
        """, unsafe_allow_html=True)

        st.session_state["intro_shown"] = True
        time.sleep(4)
        st.rerun()

    # ===== PREMIUM SPLIT-SCREEN LOGIN =====
    import os as _os, base64 as _b64_l
    _logo_file = None
    for _f in ("logo.png", "logo.jpg", "logo.jpeg", "logo.webp"):
        if _os.path.exists(_f):
            _logo_file = _f
            break
    _logo_img = ""
    if _logo_file:
        try:
            with open(_logo_file, "rb") as _lf:
                _lb64 = _b64_l.b64encode(_lf.read()).decode()
            _lext = _logo_file.rsplit(".", 1)[-1]
            _logo_img = f"<img src='data:image/{_lext};base64,{_lb64}' class='glow-logo'/>"
        except Exception:
            _logo_img = ""

    st.markdown(f"""
    <style>
        /* ===== Dark premium canvas ===== */
        .stApp {{
            background: linear-gradient(-45deg, #060d1c, #0a1f44, #131a3d, #1a1140);
            background-size: 400% 400%;
            animation: bgShift 18s ease infinite;
        }}
        @keyframes bgShift {{
            0%   {{ background-position: 0% 50%; }}
            50%  {{ background-position: 100% 50%; }}
            100% {{ background-position: 0% 50%; }}
        }}
        [data-testid="stHeader"] {{ background: transparent; }}
        .block-container {{ padding-top: 2rem !important; max-width: 1240px !important; }}

        /* ===== Floating background icons ===== */
        .float-layer {{ position: fixed; inset: 0; pointer-events: none; z-index: 0; overflow: hidden; }}
        .float-layer span {{
            position: absolute; font-size: 30px; opacity: .10;
            animation: drift linear infinite;
        }}
        @keyframes drift {{
            0%   {{ transform: translateY(105vh) rotate(0deg); }}
            100% {{ transform: translateY(-15vh) rotate(360deg); }}
        }}

        /* ===== Left brand panel ===== */
        .glow-logo {{
            width: 132px; height: 132px; border-radius: 26px; display: block;
            box-shadow: 0 0 46px rgba(124,58,237,.55), 0 0 90px rgba(37,99,235,.30);
            animation: logoPulse 4s ease-in-out infinite;
        }}
        @keyframes logoPulse {{
            0%,100% {{ transform: translateY(0); box-shadow: 0 0 46px rgba(124,58,237,.55), 0 0 90px rgba(37,99,235,.30); }}
            50%     {{ transform: translateY(-7px); box-shadow: 0 0 62px rgba(249,115,22,.45), 0 0 110px rgba(124,58,237,.40); }}
        }}
        .hero-title {{
            font-family:'Sora',sans-serif; font-size: 40px; line-height: 1.12; font-weight: 800;
            letter-spacing: -1px; color: #fff; margin: 26px 0 12px;
        }}
        .hero-title .g {{
            background: linear-gradient(110deg,#3b82f6,#a855f7,#f97316);
            -webkit-background-clip: text; background-clip: text; color: transparent;
        }}
        .hero-sub {{ color:#a9b5d1; font-size: 16.5px; line-height:1.6; margin-bottom: 6px; }}
        .hero-sub b {{ color:#e6c866; }}

        /* ===== Rotating feature cards ===== */
        .rot-wrap {{ position: relative; height: 92px; margin-top: 26px; }}
        .rot-card {{
            position: absolute; inset: 0; opacity: 0;
            background: rgba(255,255,255,.06); border: 1px solid rgba(255,255,255,.13);
            backdrop-filter: blur(14px); border-radius: 16px;
            padding: 16px 20px; display: flex; align-items: center; gap: 15px;
            animation: rotate8 32s infinite;
        }}
        .rot-card .ric {{
            width: 46px; height: 46px; border-radius: 13px; flex-shrink: 0;
            background: linear-gradient(135deg, rgba(59,130,246,.30), rgba(168,85,247,.20));
            display: flex; align-items: center; justify-content: center; font-size: 22px;
        }}
        .rot-card b {{ color:#fff; font-family:'Sora',sans-serif; font-size:15.5px; display:block; }}
        .rot-card span {{ color:#a9b5d1; font-size:13px; }}
        .rot-card:nth-child(1) {{ animation-delay: 0s; }}
        .rot-card:nth-child(2) {{ animation-delay: 4s; }}
        .rot-card:nth-child(3) {{ animation-delay: 8s; }}
        .rot-card:nth-child(4) {{ animation-delay: 12s; }}
        .rot-card:nth-child(5) {{ animation-delay: 16s; }}
        .rot-card:nth-child(6) {{ animation-delay: 20s; }}
        .rot-card:nth-child(7) {{ animation-delay: 24s; }}
        .rot-card:nth-child(8) {{ animation-delay: 28s; }}
        @keyframes rotate8 {{
            0%    {{ opacity: 0; transform: translateY(12px); }}
            2%    {{ opacity: 1; transform: translateY(0); }}
            11%   {{ opacity: 1; transform: translateY(0); }}
            13%   {{ opacity: 0; transform: translateY(-12px); }}
            100%  {{ opacity: 0; transform: translateY(-12px); }}
        }}

        /* ===== Glass login card (styles Streamlit's bordered container) ===== */
        [data-testid="stVerticalBlockBorderWrapper"] {{
            background: rgba(255,255,255,.055) !important;
            border: 1px solid rgba(255,255,255,.15) !important;
            border-radius: 24px !important;
            backdrop-filter: blur(22px);
            box-shadow: 0 26px 70px rgba(0,0,0,.42);
            padding: 10px 8px !important;
        }}
        .glass-title {{ font-family:'Sora',sans-serif; font-size:24px; font-weight:800; color:#fff; text-align:center; }}
        .glass-sub {{ color:#a9b5d1; font-size:13.5px; text-align:center; margin-bottom:4px; }}
        .trial-chip {{
            background: linear-gradient(135deg, rgba(249,115,22,.20), rgba(168,85,247,.18));
            border: 1px solid rgba(249,115,22,.45); border-radius: 12px;
            padding: 11px 14px; text-align: center; margin: 6px 0 4px;
        }}
        .trial-chip b {{ color:#fbbf24; font-size:14.5px; }}
        .trial-chip span {{ color:#d8b4fe; font-size:12.5px; }}

        /* Google button premium styling */
        div[data-testid="stVerticalBlockBorderWrapper"] a,
        div[data-testid="stVerticalBlockBorderWrapper"] button {{
            border-radius: 13px !important; font-weight: 700 !important;
        }}

        .trust-row {{ display:flex; justify-content:center; gap:7px; flex-wrap:wrap; margin-top:12px; }}
        .trust-row span {{
            font-size:11.5px; color:#c7d1e6; background:rgba(255,255,255,.07);
            border:1px solid rgba(255,255,255,.13); padding:5px 11px; border-radius:20px;
        }}

        /* ===== Footer ===== */
        .prem-foot {{
            position: relative; z-index: 2; text-align:center; color:#8a97b5; font-size:12.5px;
            margin-top:44px; padding-top:22px; border-top:1px solid rgba(255,255,255,.10); line-height:1.8;
        }}
        .prem-foot .gst {{
            display:inline-block; background:rgba(168,85,247,.14); border:1px solid rgba(168,85,247,.35);
            color:#d8b4fe; padding:3px 11px; border-radius:6px; font-family:monospace; font-size:12px;
        }}
        .prem-foot a {{ color:#60a5fa; text-decoration:none; font-weight:600; }}

        @media (max-width: 900px) {{
            .hero-title {{ font-size: 30px; }}
            .glow-logo {{ width: 100px; height: 100px; }}
        }}
    </style>

    <div class="float-layer">
        <span style="left:6%;  animation-duration:26s; animation-delay:0s;">📘</span>
        <span style="left:19%; animation-duration:34s; animation-delay:4s;">🧠</span>
        <span style="left:33%; animation-duration:29s; animation-delay:9s;">✏️</span>
        <span style="left:47%; animation-duration:38s; animation-delay:2s;">📄</span>
        <span style="left:61%; animation-duration:31s; animation-delay:12s;">🎓</span>
        <span style="left:74%; animation-duration:27s; animation-delay:6s;">⭐</span>
        <span style="left:88%; animation-duration:36s; animation-delay:15s;">📊</span>
        <span style="left:95%; animation-duration:30s; animation-delay:10s;">🤖</span>
    </div>
    """, unsafe_allow_html=True)

    left_col, right_col = st.columns([1.18, 1], gap="large")

    with left_col:
        st.markdown(f"""
        <div style="position:relative; z-index:2; padding-top:6px;">
            {_logo_img}
            <div class="hero-title">AI Powered<br><span class="g">Question Paper Generator</span></div>
            <div class="hero-sub">Generate professional question papers within <b>3 minutes</b> — built on the TN State Board pattern.</div>
            <div class="rot-wrap">
                <div class="rot-card"><div class="ric">⚡</div><div><b>Ready in 3 Minutes</b><span>From blueprint to Word file</span></div></div>
                <div class="rot-card"><div class="ric">📚</div><div><b>Classes 6 – 12</b><span>All subjects</span></div></div>
                <div class="rot-card"><div class="ric">🎯</div><div><b>TN State Board Pattern</b><span>Blueprint based</span></div></div>
                <div class="rot-card"><div class="ric">📄</div><div><b>Export to Word</b><span>One click download</span></div></div>
                <div class="rot-card"><div class="ric">🤖</div><div><b>AI Difficulty Levels</b><span>Easy · Medium · Hard</span></div></div>
                <div class="rot-card"><div class="ric">📝</div><div><b>25 / 50 / 75 / 100 Marks</b><span>Any exam size</span></div></div>
                <div class="rot-card"><div class="ric">📖</div><div><b>Book Back · Exercise · Example</b><span>Real textbook question bank</span></div></div>
                <div class="rot-card"><div class="ric">📊</div><div><b>Auto Blueprint & Choice</b><span>Sections handled for you</span></div></div>
            </div>
        </div>
        """, unsafe_allow_html=True)

    with right_col:
        _card = st.container(border=True)
        with _card:
            st.markdown("""
            <div style="padding:14px 6px 2px;">
                <div class="glass-title">Welcome 👋</div>
                <div class="glass-sub">Sign in to start creating question papers</div>
                <div class="trial-chip">
                    <b>🎁 1 Month Free Trial</b><br>
                    <span>Unlimited question papers · No card needed</span>
                </div>
            </div>
            """, unsafe_allow_html=True)

            result = oauth2.authorize_button(
                name="Continue with Google",
                icon="https://www.google.com/favicon.ico",
                redirect_uri=GOOGLE_REDIRECT_URI,
                scope="openid email profile",
                key="google_login",
                extras_params={"prompt": "select_account"},
                use_container_width=True,
            )

            st.markdown("""
            <div class="trust-row">
                <span>✅ AI Powered</span>
                <span>🔒 Secure Login</span>
                <span>☁️ Cloud Based</span>
                <span>🔄 Always Updated</span>
            </div>
            <div style="height:8px;"></div>
            """, unsafe_allow_html=True)

    st.markdown("""
    <div class="prem-foot">
        <span class="gst">GSTIN: 33ABJFP1752G1ZC</span><br><br>
        <a href="https://www.facebook.com/profile.php?id=61590340754238" target="_blank">📘 Follow us on Facebook</a><br>
        📍 39 to 41, Gayathri Complex, Kumarasamypatty Cherry Road, Hasthampatty, Salem – 636007<br>
        📞 +91 90430 00733 &nbsp;·&nbsp; © 2026 PMP Enterprises · PMP Edu AI
    </div>
    """, unsafe_allow_html=True)

    if result and "token" in result:
        token = result["token"]
        access_token = token.get("access_token", "")
        guser = _google_userinfo(access_token)
        if guser and "email" in guser:
            db_user = upsert_google_user(guser)
            if not db_user:
                # Fallback: fetch existing row by email so we still get real id + profile
                db_user = fetch_user_by_email(guser["email"])
            st.session_state["logged_in_user"] = db_user or {
                "id":      None,
                "email":   guser.get("email", ""),
                "name":    guser.get("name", "User"),
                "picture": guser.get("picture", ""),
                "plan":    "free",
            }
            st.rerun()

    st.stop()


# ==========================================
# LOGGED IN — Get user & check usage
# ==========================================
current_user = st.session_state["logged_in_user"]

# Always refresh from DB so profile fields (school/mobile) are accurate and
# the setup gate never re-triggers for users who already completed it.
_fresh = fetch_user_by_email(current_user.get("email", "")) if current_user.get("email") else None
if _fresh:
    current_user = _fresh
    st.session_state["logged_in_user"] = current_user

user_id   = current_user["id"]
user_name = current_user["name"]
user_email= current_user["email"]
user_plan = current_user.get("plan", "free")
user_pic  = current_user.get("picture", "")

# ==========================================
# PROFILE SETUP GATE — first login only (skip if DB read failed)
# ==========================================
_needs_profile = (_fresh is not None) and not (current_user.get("mobile") and current_user.get("school_name"))
if _needs_profile:
    st.markdown("""
    <div style='text-align:center; padding:30px 0 10px 0;'>
        <div style='font-size:52px;'>👋</div>
        <h1 style='color:#0a1f44; margin:8px 0 4px 0;'>வரவேற்கிறோம்!</h1>
        <p style='color:#64748b; font-size:16px;'>தொடங்கும் முன், உங்கள் விவரங்களை பூர்த்தி செய்யவும்</p>
    </div>
    """, unsafe_allow_html=True)

    pcol1, pcol2, pcol3 = st.columns([1, 2, 1])
    with pcol2:
        with st.form("profile_setup_form"):
            st.markdown("#### 🏫 Profile உருவாக்கம்")
            in_school  = st.text_input("பள்ளியின் பெயர் *", placeholder="உ.தா: அரசு மேல்நிலைப் பள்ளி, சேலம்")
            in_teacher = st.text_input("ஆசிரியர் பெயர் *", value=user_name, placeholder="உங்கள் முழு பெயர்")
            in_mobile  = st.text_input("மொபைல் எண் *", placeholder="10 இலக்க எண்", max_chars=10)
            st.caption("* அனைத்து புலங்களும் அவசியம். இது ஒரு முறை மட்டுமே கேட்கப்படும்.")
            submitted = st.form_submit_button("✅ சேமித்து தொடங்கு", use_container_width=True, type="primary")

            if submitted:
                mobile_clean = in_mobile.strip()
                if not in_school.strip() or not in_teacher.strip() or not mobile_clean:
                    st.error("⚠️ அனைத்து புலங்களையும் நிரப்பவும்.")
                elif not (mobile_clean.isdigit() and len(mobile_clean) == 10):
                    st.error("⚠️ சரியான 10 இலக்க மொபைல் எண்ணை உள்ளிடவும்.")
                else:
                    if update_user_profile(user_id, in_school, in_teacher, mobile_clean, email=user_email):
                        current_user["school_name"]  = in_school.strip()
                        current_user["teacher_name"] = in_teacher.strip()
                        current_user["mobile"]       = mobile_clean
                        st.session_state["logged_in_user"] = current_user
                        st.success("✅ Profile உருவாக்கப்பட்டது!")
                        st.rerun()
    st.stop()

# ==========================================
# LOGGED IN — check trial period
# ==========================================
user_school  = current_user.get("school_name", "")
user_teacher = current_user.get("teacher_name", user_name)

TRIAL_DAYS = 30
is_premium = user_plan in ["premium", "paid"]

# Compute trial status from created_at (first login date)
import datetime as _dt
_created = current_user.get("created_at")
trial_days_left = None
trial_expired = False
if not is_premium and _created:
    try:
        if isinstance(_created, str):
            _created_dt = _dt.datetime.fromisoformat(_created.replace("Z", "").split(".")[0])
        else:
            _created_dt = _created
        _elapsed = (_dt.datetime.now() - _created_dt.replace(tzinfo=None)).days
        trial_days_left = max(TRIAL_DAYS - _elapsed, 0)
        trial_expired = _elapsed >= TRIAL_DAYS
    except Exception:
        trial_days_left = TRIAL_DAYS
        trial_expired = False

# Block completely when trial expired (and not premium)
if trial_expired and not is_premium:
    st.markdown("""
    <div style='text-align:center; padding:50px 20px;'>
        <div style='font-size:56px;'>⏳</div>
        <h1 style='color:#0a1f44;'>உங்கள் 30-நாள் இலவச சோதனை முடிந்தது</h1>
        <p style='color:#64748b; font-size:16px; margin-top:8px;'>
            தொடர்ந்து வினாத்தாள்கள் உருவாக்க, Premium plan-க்கு மேம்படுத்தவும்.
        </p>
        <div style='margin-top:20px; padding:16px 22px; background:#fff6df; border:1px solid #c9a227;
                    border-radius:12px; display:inline-block;'>
            <b style='color:#9a6b00;'>📞 தொடர்பு:</b> +91 90430 00733<br>
            <span style='color:#9a6b00; font-size:14px;'>PMP Enterprises · Premium plan விவரங்களுக்கு அழைக்கவும்</span>
        </div>
    </div>
    """, unsafe_allow_html=True)
    _lo1, _lo2, _lo3 = st.columns([2, 1, 2])
    with _lo2:
        if st.button("🚪 Logout", use_container_width=True):
            st.session_state.pop("logged_in_user", None); st.session_state.pop("intro_shown", None)
            st.rerun()
    st.stop()

# ==========================================
# TRIAL BANNER
# ==========================================
if is_premium:
    st.success(f"⭐ Premium Plan | 👤 {user_name} | Unlimited Access")
elif trial_days_left is not None:
    if trial_days_left <= 5:
        st.warning(f"⏳ இலவச சோதனை: இன்னும் {trial_days_left} நாட்கள் மட்டுமே மிச்சம் | வரம்பற்ற வினாத்தாள்கள் | 👤 {user_name}")
    else:
        st.info(f"🎁 இலவச சோதனை காலம்: இன்னும் {trial_days_left} நாட்கள் · வரம்பற்ற வினாத்தாள்கள் 🚀 | 👤 {user_name}")

# ==========================================
# API Configuration
# ==========================================
YOUR_API_KEY = st.secrets["GEMINI_API_KEY"]
client = genai.Client(api_key=YOUR_API_KEY)

def gemini_generate(prompt, model='gemini-2.5-flash', max_retries=3):
    """Gemini call with automatic 429/503 retry + backoff. Returns response or raises."""
    last_err = None
    for attempt in range(max_retries):
        try:
            return client.models.generate_content(model=model, contents=prompt)
        except Exception as e:
            last_err = e
            msg = str(e)
            if "429" in msg or "RESOURCE_EXHAUSTED" in msg or "503" in msg:
                wait = 6 * (attempt + 1)
                m = re.search(r'retry in (\d+)', msg, re.IGNORECASE) or re.search(r"retryDelay['\":\s]+(\d+)", msg)
                if m:
                    wait = min(int(m.group(1)) + 1, 30)
                time.sleep(wait)
            else:
                raise
    raise last_err

# ==========================================
# Database Loading
# ==========================================
@st.cache_data
def load_data():
    try:
        return pd.read_csv('lesson_master_v1_5.csv')
    except:
        return pd.DataFrame()

ensure_question_bank_table()

def get_math_dynamic_weightage(selected_lessons, part1_val, part2_val, part3_val):
    base_matrix = {
        "Relations and Functions":      {"1M": 1.5, "2M": 2, "5M": 1.5, "8M": 0},
        "Numbers and Sequences":        {"1M": 2.0, "2M": 2, "5M": 2.0, "8M": 0},
        "Algebra":                      {"1M": 2.0, "2M": 2, "5M": 2.0, "8M": 1},
        "Geometry":                     {"1M": 2.0, "2M": 1, "5M": 1.0, "8M": 1},
        "Coordinate Geometry":          {"1M": 1.5, "2M": 2, "5M": 2.0, "8M": 0},
        "Mensuration":                  {"1M": 1.5, "2M": 2, "5M": 2.0, "8M": 0},
        "Statistics and Probability":   {"1M": 2.0, "2M": 2, "5M": 2.0, "8M": 0}
    }
    rules = []
    for lesson in selected_lessons:
        if lesson in base_matrix:
            bm = base_matrix[lesson]
            rules.append(f"- From '{lesson}': Generate approx {int(bm['1M'])} MCQs, {bm['2M']} Questions (2-Mark), {int(bm['5M'])} Questions (5-Mark), and {bm['8M']} Question (8-Mark).")
    return "\n".join(rules)

def get_english_blueprint_rules():
    return """
    [STRICT MASTER ENGLISH BLUEPRINT LOCK]
    PART I (14 Marks): One Mark Questions
    - Q1-3: Synonyms strictly selected from the text prose context.
    - Q4-6: Antonyms strictly selected from the text prose context.
    - Q7-14: Textual Grammar Matrix: Plural Form, Suffix/Prefix, Abbreviations, Phrasal Verbs, Compound Words, Prepositions, Tense Forms, Linkers.
    PART II (20 Marks): Two Mark Questions (Answer any 10 out of 12)
    - Section 1 (Prose Qs): 3 Short answer questions from selected textbook lessons.
    - Section 2 (Poetry Appreciation): 3 Poetic line sets with internal appreciation questions.
    - Section 3 (Grammar Blocks): 5 Mandatory Core Grammar Questions: Voice Change, Reported Speech, Punctuation, Sentence Transformation, Rearrange jumbled words.
    - Section 4 (Roadmap): 1 Question on Road Map directions guide.
    PART III (35 Marks): Five Mark Questions (Answer any 7 out of 10)
    - Section 1 (Literature Paragraphs): Prose Paragraphs and Poetry Paragraphs covering textbook lessons.
    - Section 2 (Coherent Order & Comprehension): Rearrange sentences in coherent order, Supplementary passage comprehension.
    - Section 3 (Writing Skills Matrix): Advertisement, Formal Letter, Notice Writing, Picture Comprehension.
    PART IV (16 Marks): Eight Mark Questions (Answer both - Internal Choice)
    - Q37: Comprehensive Prose Passage reading or Poem reading with granularity questions.
    - Q38: Detailed Literature Essay from supplementary stories.
    """

def get_blueprint_defaults(total_marks, is_social=False, is_english=False):
    if is_english or is_social:
        return {"p1": 14, "p2g": 12, "p2a": 10, "p3g": 10, "p3a": 7, "p4v": 8, "p4g": 2, "p4a": 2}
    defaults = {"p1": 14, "p2g": 12, "p2a": 10, "p3g": 14, "p3a": 10, "p4v": 8, "p4g": 4, "p4a": 2}
    if total_marks == 106:
        defaults = {"p1": 20, "p2g": 12, "p2a": 10, "p3g": 14, "p3a": 10, "p4v": 8, "p4g": 4, "p4a": 2}
    elif total_marks == 50:
        defaults = {"p1": 10, "p2g": 8, "p2a": 6, "p3g": 6, "p3a": 4, "p4v": 8, "p4g": 2, "p4a": 1}
    elif total_marks == 25:
        defaults = {"p1": 5, "p2g": 6, "p2a": 5, "p3g": 3, "p3a": 2, "p4v": 8, "p4g": 0, "p4a": 0}
    return defaults

def _clean_diagram_label(label_text):
    """matplotlib Tamil render பண்ணாது - ASCII மட்டும் வைக்கிறோம்"""
    import unicodedata
    cleaned = label_text.replace("Angle", "∠").replace("angle", "∠")
    # Keep only ASCII + common math symbols (Tamil letters become boxes in matplotlib)
    safe = ''.join(c for c in cleaned if ord(c) < 0x0B80 or ord(c) > 0x0BFF)
    return safe.strip()

def generate_geometry_image(shape_type, label_text=""):
    fig, ax = plt.subplots(figsize=(2.8, 2.8))
    shape_upper = shape_type.upper()
    clean_label = _clean_diagram_label(label_text)

    if "THALES" in shape_upper or "BPT" in shape_upper:
        # தேல்ஸ் தேற்றம் / Basic Proportionality Theorem
        # Triangle ABC with D on AB, E on AC, DE parallel to BC
        A = np.array([2.0, 3.6]); B = np.array([0.2, 0.2]); C = np.array([3.8, 0.2])
        # D and E at 45% down from A
        t = 0.45
        D = A + t * (B - A)
        E = A + t * (C - A)
        # Main triangle
        tri = np.array([A, B, C, A])
        ax.plot(tri[:, 0], tri[:, 1], 'k-', lw=2)
        # DE parallel line
        ax.plot([D[0], E[0]], [D[1], E[1]], 'k-', lw=2)
        # Parallel marks on DE and BC (small double ticks)
        for seg_p1, seg_p2 in [(D, E), (B, C)]:
            mid = (seg_p1 + seg_p2) / 2
            ax.annotate('▸', xy=mid, fontsize=8, ha='center', va='center')
        # Labels
        ax.text(A[0], A[1]+0.15, 'A', fontsize=12, fontweight='bold', ha='center')
        ax.text(B[0]-0.2, B[1]-0.1, 'B', fontsize=12, fontweight='bold')
        ax.text(C[0]+0.1, C[1]-0.1, 'C', fontsize=12, fontweight='bold')
        ax.text(D[0]-0.25, D[1], 'D', fontsize=12, fontweight='bold')
        ax.text(E[0]+0.12, E[1], 'E', fontsize=12, fontweight='bold')
        # DE ∥ BC annotation
        ax.text(2.0, -0.35, 'DE ∥ BC', fontsize=11, ha='center', fontweight='bold')
        ax.set_xlim(-0.5, 4.4); ax.set_ylim(-0.7, 4.1)

    elif "EXT_BISECTOR" in shape_upper or "EXTERNAL_BISECTOR" in shape_upper:
        # வெளிப்புற கோண இருசமவெட்டி - AD meets BC EXTENSION at D
        A = np.array([1.4, 3.0]); B = np.array([0.3, 0.5]); C = np.array([2.9, 0.5])
        D = np.array([4.6, 0.5])  # on BC extension
        tri = np.array([A, B, C, A])
        ax.plot(tri[:, 0], tri[:, 1], 'k-', lw=2)
        # BC extension (dashed)
        ax.plot([C[0], D[0]], [C[1], D[1]], 'k--', lw=1.5)
        # AD line
        ax.plot([A[0], D[0]], [A[1], D[1]], 'k-', lw=1.8)
        ax.plot(*D, 'ko', markersize=4)
        ax.text(A[0], A[1]+0.15, 'A', fontsize=12, fontweight='bold', ha='center')
        ax.text(B[0]-0.25, B[1]-0.15, 'B', fontsize=12, fontweight='bold')
        ax.text(C[0]-0.1, C[1]-0.35, 'C', fontsize=12, fontweight='bold')
        ax.text(D[0]+0.08, D[1]-0.15, 'D', fontsize=12, fontweight='bold')
        ax.set_xlim(-0.3, 5.3); ax.set_ylim(-0.4, 3.6)

    elif "ANGLE_BISECTOR" in shape_upper or "BISECTOR" in shape_upper:
        # உட்புற கோண இருசமவெட்டி - AD bisects angle A, D on BC
        A = np.array([2.0, 3.4]); B = np.array([0.3, 0.4]); C = np.array([3.9, 0.4])
        D = np.array([2.35, 0.4])  # on BC (bisector foot, slightly right of midpoint)
        tri = np.array([A, B, C, A])
        ax.plot(tri[:, 0], tri[:, 1], 'k-', lw=2)
        ax.plot([A[0], D[0]], [A[1], D[1]], 'k-', lw=1.8)
        ax.plot(*D, 'ko', markersize=4)
        # angle bisector arcs at A (two small equal-angle marks)
        ax.annotate('', xy=(1.75, 2.85), xytext=(1.9, 3.0),
                    arrowprops=dict(arrowstyle='-', lw=1))
        ax.annotate('', xy=(2.25, 2.85), xytext=(2.1, 3.0),
                    arrowprops=dict(arrowstyle='-', lw=1))
        ax.text(A[0], A[1]+0.15, 'A', fontsize=12, fontweight='bold', ha='center')
        ax.text(B[0]-0.25, B[1]-0.15, 'B', fontsize=12, fontweight='bold')
        ax.text(C[0]+0.1, C[1]-0.15, 'C', fontsize=12, fontweight='bold')
        ax.text(D[0]+0.05, D[1]-0.35, 'D', fontsize=12, fontweight='bold')
        ax.set_xlim(-0.3, 4.5); ax.set_ylim(-0.5, 4.0)

    elif "TWO_TANGENT" in shape_upper or "TANGENTS" in shape_upper:
        # வெளிப்புள்ளியில் இருந்து இரண்டு தொடுகோடுகள் PA, PB
        O = np.array([1.5, 2.0]); r = 1.1
        P = np.array([4.3, 2.0])
        circle = plt.Circle(O, r, fill=False, color='black', lw=2)
        ax.add_patch(circle)
        ax.plot(*O, 'ko', markersize=4)
        ax.text(O[0]-0.05, O[1]+0.15, 'O', fontsize=11, fontweight='bold')
        # Tangent points A (top) and B (bottom)
        d = np.linalg.norm(P - O)
        ang = np.arccos(r / d)
        base = np.arctan2(P[1]-O[1], P[0]-O[0])
        Apt = O + r * np.array([np.cos(base + ang), np.sin(base + ang)])
        Bpt = O + r * np.array([np.cos(base - ang), np.sin(base - ang)])
        ax.plot([P[0], Apt[0]], [P[1], Apt[1]], 'k-', lw=1.8)
        ax.plot([P[0], Bpt[0]], [P[1], Bpt[1]], 'k-', lw=1.8)
        # Radii OA, OB (dashed)
        ax.plot([O[0], Apt[0]], [O[1], Apt[1]], 'k--', lw=1)
        ax.plot([O[0], Bpt[0]], [O[1], Bpt[1]], 'k--', lw=1)
        ax.plot(*Apt, 'ko', markersize=4)
        ax.plot(*Bpt, 'ko', markersize=4)
        ax.plot(*P, 'ko', markersize=4)
        ax.text(Apt[0]-0.1, Apt[1]+0.18, 'A', fontsize=12, fontweight='bold')
        ax.text(Bpt[0]-0.1, Bpt[1]-0.32, 'B', fontsize=12, fontweight='bold')
        ax.text(P[0]+0.1, P[1]-0.05, 'P', fontsize=12, fontweight='bold')
        ax.set_xlim(0, 5.2); ax.set_ylim(0.2, 3.8)
        ax.set_aspect('equal')

    elif "PYTHAGORAS" in shape_upper or "RIGHT" in shape_upper:
        # செங்கோண முக்கோணம் - right angle at B
        A = np.array([0.3, 3.4]); B = np.array([0.3, 0.3]); C = np.array([3.9, 0.3])
        tri = np.array([A, B, C, A])
        ax.plot(tri[:, 0], tri[:, 1], 'k-', lw=2)
        # Right angle square at B
        ax.plot([0.3, 0.7], [0.7, 0.7], 'k-', lw=1)
        ax.plot([0.7, 0.7], [0.3, 0.7], 'k-', lw=1)
        ax.text(A[0]-0.05, A[1]+0.15, 'A', fontsize=12, fontweight='bold', ha='center')
        ax.text(B[0]-0.25, B[1]-0.15, 'B', fontsize=12, fontweight='bold')
        ax.text(C[0]+0.1, C[1]-0.15, 'C', fontsize=12, fontweight='bold')
        ax.set_xlim(-0.4, 4.4); ax.set_ylim(-0.6, 4.0)

    elif "TANGENT" in shape_upper:
        # வட்டத்துக்கு தொடுகோடு
        circle = plt.Circle((1.8, 2.0), 1.2, fill=False, color='black', lw=2)
        ax.add_patch(circle)
        ax.plot(1.8, 2.0, 'ko', markersize=4)
        ax.text(1.85, 2.1, 'O', fontsize=11, fontweight='bold')
        # Tangent point P at right side of circle
        P = np.array([3.0, 2.0])
        ax.plot(*P, 'ko', markersize=4)
        ax.text(P[0]+0.1, P[1]+0.1, 'P', fontsize=11, fontweight='bold')
        # Radius OP
        ax.plot([1.8, P[0]], [2.0, P[1]], 'k-', lw=1.2)
        # Tangent line (vertical through P)
        ax.plot([P[0], P[0]], [0.4, 3.6], 'k-', lw=2)
        # Right angle mark at P
        ax.plot([P[0]-0.22, P[0]-0.22], [2.0, 2.22], 'k-', lw=1)
        ax.plot([P[0]-0.22, P[0]], [2.22, 2.22], 'k-', lw=1)
        ax.set_xlim(0, 4.2); ax.set_ylim(0, 4.0)

    elif "TRIANGLE" in shape_upper:
        points = np.array([[0.3, 0.3], [4.1, 0.3], [2.2, 3.3], [0.3, 0.3]])
        ax.plot(points[:, 0], points[:, 1], 'k-', lw=2)
        ax.text(0.1, 0.05, 'A', fontsize=12, fontweight='bold')
        ax.text(4.2, 0.05, 'B', fontsize=12, fontweight='bold')
        ax.text(2.2, 3.45, 'C', fontsize=12, fontweight='bold', ha='center')
        ax.set_xlim(-0.3, 4.7); ax.set_ylim(-0.6, 3.9)

    elif "SQUARE" in shape_upper or "RECTANGLE" in shape_upper:
        w = 4 if "RECTANGLE" in shape_upper else 3
        points = np.array([[0.3, 0.3], [w+0.3, 0.3], [w+0.3, 3.3], [0.3, 3.3], [0.3, 0.3]])
        ax.plot(points[:, 0], points[:, 1], 'k-', lw=2)
        ax.text(0.05, 0.05, 'A', fontsize=11, fontweight='bold')
        ax.text(w+0.4, 0.05, 'B', fontsize=11, fontweight='bold')
        ax.text(w+0.4, 3.4, 'C', fontsize=11, fontweight='bold')
        ax.text(0.05, 3.4, 'D', fontsize=11, fontweight='bold')
        ax.set_xlim(-0.4, w+1); ax.set_ylim(-0.6, 3.9)

    elif "SEMICIRCLE" in shape_upper or "SEMI" in shape_upper:
        theta = np.linspace(0, np.pi, 100)
        ax.plot(2 + 1.6*np.cos(theta), 1 + 1.6*np.sin(theta), 'k-', lw=2)
        ax.plot([0.4, 3.6], [1, 1], 'k-', lw=2)
        ax.plot(2, 1, 'ko', markersize=4)
        ax.text(2.05, 0.78, 'O', fontsize=11, fontweight='bold')
        ax.set_xlim(0, 4); ax.set_ylim(0, 3.2)

    elif "CIRCLE" in shape_upper:
        circle = plt.Circle((2, 2), 1.6, fill=False, color='black', lw=2)
        ax.add_patch(circle)
        ax.plot(2, 2, 'ko', markersize=4)
        ax.text(2.1, 2.1, 'O', fontsize=11, fontweight='bold')
        ax.plot([2, 3.6], [2, 2], 'k-', lw=1.2)
        ax.text(2.8, 2.14, 'r', fontsize=10, style='italic')
        ax.set_xlim(0, 4); ax.set_ylim(0, 4)

    # Label (ASCII-safe only)
    if clean_label:
        ax.text(0.5, 0.01, clean_label, fontsize=9, ha='center',
                fontweight='bold', color='blue', transform=ax.transAxes)

    ax.set_aspect('equal')
    ax.axis('off')
    img_buf = io.BytesIO()
    plt.savefig(img_buf, format='png', bbox_inches='tight', dpi=110)
    img_buf.seek(0)
    plt.close(fig)
    return img_buf

def generate_prompt_v18(subject, lessons_list, exam_type, exam_time, total_marks, exam_mode, blueprint_desc, part1_val, part2_val, part3_val, diff_level, paper_lang="தமிழ் (Tamil)"):
    lessons_str = ", ".join(lessons_list)
    sub_lower = subject.lower()
    is_english = "english" in sub_lower or "ஆங்கிலம்" in sub_lower
    is_tamil   = "tamil"   in sub_lower or "தமிழ்"    in sub_lower
    is_social  = "social"  in sub_lower or "சமூக"     in sub_lower
    is_math    = "math"    in sub_lower or "கணிதம்"   in sub_lower

    # User-chosen paper language overrides the auto default (except for language subjects
    # Tamil/English where the medium is fixed by the subject itself).
    force_english = paper_lang.startswith("English") and not is_tamil
    force_tamil   = paper_lang.startswith("தமிழ்") and not is_english

    if diff_level == "எளிமை (Easy)":
        difficulty_directive = "DIFFICULTY CRITERIA: Focus 80% on direct textbook back questions and formulas."
    elif diff_level == "நடுத்தரம் (Medium)":
        difficulty_directive = "DIFFICULTY CRITERIA: Balanced public paper structure. 60% Direct, 30% Application, 10% HOTS."
    else:
        difficulty_directive = "DIFFICULTY CRITERIA: High-level revision standard with indirect grammatical transformations."

    math_weightage_directive = ""
    if is_math:
        math_weightage_directive = f"[STRICT LESSON-WISE MARKS WEIGHTAGE MATRIX]\n{get_math_dynamic_weightage(lessons_list, part1_val, part2_val, part3_val)}"

    if is_english:
        lang_instruction      = "5. Language: Pure ENGLISH only."
        header_format         = "PART [ROMAN_NUM] - [Section Description] (No_of_Qs x Marks = Total_Marks)"
        option_format         = "Options marker: a) , b) , c) , d)"
        subject_blueprint_rules = f"[STRICT TN BOARD ENGLISH BLUEPRINT LOCK]\n{get_english_blueprint_rules()}"
    elif is_tamil:
        lang_instruction      = "5. Language: Pure TAMIL only."
        header_format         = "பகுதி [ROMAN_NUM] - [பிரிவின் விளக்கம்] (வினாக்கள் எண்ணிக்கை x மதிப்பெண் = மொத்த மதிப்பெண்கள்)"
        option_format         = "Options marker: அ) , ஆ) , இ) , ஈ)"
        subject_blueprint_rules = "[அசல் தமிழ் பாடத்திட்ட ப்ளூபிரின்ட்] சொல்வளம், இலக்கணம் லாக்."
    elif is_social:
        lang_instruction      = "5. Language: Pure TAMIL only."
        header_format         = "பகுதி [ROMAN_NUM] - [பிரிவின் விளக்கம்]"
        option_format         = "Options marker: அ) , ஆ) , இ) , ஈ)"
        subject_blueprint_rules = "[MANDATORY CRITICAL SOCIAL SCIENCE BLUEPRINT] Assertion-Reason, Map locked."
    elif is_math:
        lang_instruction      = "5. Language: Pure TAMIL only."
        header_format         = "பகுதி [ROMAN_NUM] - [பிரிவின் விளக்கம்] (No_of_Qs x Marks = Total_Marks)"
        option_format         = "Options marker: அ) , ஆ) , இ) , ஈ)"
        subject_blueprint_rules = f"""[MANDATORY CRITICAL MATHEMATICS CORE EMBEDDED LOCK]
1. ABSOLUTE BAN ON AI DISCLAIMERS.
2. GEOMETRY DIAGRAM TAGS - STRICT RULES:
   - NEVER draw diagrams using text characters (/, \\, -, |). ASCII art is ABSOLUTELY BANNED.
   - Instead, write the FULL question text first, then on the NEXT separate line add ONE diagram tag.
   - The question text must always be complete. NEVER replace question text with a tag.
   - TAG SELECTION MATRIX (match question content EXACTLY - wrong diagram = INVALID paper):
     * Question mentions "DE" and "BC" or "DE ∥ BC" or "மிகைவிகித"/"விகிதசம" → [DRAW_THALES] (NEVER plain TRIANGLE)
     * தேல்ஸ் தேற்றம் / Thales / BPT → [DRAW_THALES]
     * உட்புற கோண இருசமவெட்டி / internal angle bisector / "AD bisects ∠A" → [DRAW_ANGLE_BISECTOR]
     * வெளிப்புற கோண இருசமவெட்டி / external bisector / "D on BC extension"/"BC-ன் நீட்டிப்பு" → [DRAW_EXT_BISECTOR]
     * ஒரு தொடுகோடு / single tangent → [DRAW_TANGENT]
     * இரண்டு தொடுகோடுகள் / two tangents PA and PB from external point → [DRAW_TWO_TANGENTS]
     * பிதாகரஸ் தேற்றம் / Pythagoras / right-angle triangle → [DRAW_PYTHAGORAS]
   - CRITICAL: If question text mentions points D, E on triangle sides → NEVER use [DRAW_TRIANGLE].
     The diagram MUST show every point named in the question (A,B,C,D,E,P etc).
     A diagram missing a point mentioned in the question is a CRITICAL ERROR.
   - GENERIC shape tags: [DRAW_TRIANGLE], [DRAW_SQUARE], [DRAW_RECTANGLE], [DRAW_CIRCLE], [DRAW_SEMICIRCLE]
   - Labels inside tags must be ENGLISH/numbers only (e.g. [DRAW_CIRCLE: r=7cm]). NO Tamil text inside tags.
   - Example CORRECT:
     35. தேல்ஸ் தேற்றத்தை எழுதி நிரூபிக்கவும்.
     [DRAW_THALES]
3. GRAPH PAPER COORDINATES.
{math_weightage_directive}"""
    else:
        lang_instruction      = "5. Language: Pure TAMIL language only."
        header_format         = "பகுதி [ROMAN_NUM]"
        option_format         = "Options marker: அ) , ஆ) , இ) , ஈ)"
        subject_blueprint_rules = ""

    no_latex_rule = """
[STRICT NO-LATEX OUTPUT RULE - CRITICAL]
- NEVER output LaTeX syntax. This is a printable exam paper, not a LaTeX document.
- BANNED: \\rightarrow, \\frac{}{}, \\{ \\}, \\times, \\sqrt{}, x^{2}, $...$, \\in, \\triangle etc.
- USE INSTEAD plain Unicode: → for arrows, { } for sets, × for multiply,
  √ for root, x² x³ superscripts, ∈ for element-of, ∠ for angle, △ for triangle,
  a/b format for fractions, θ π α β for Greek letters.
- Example CORRECT: f: A → B, f(x) = 3x - 1, A = {1, 2, 3, 4}, x² + y² = r²
- Example WRONG: f: A \\rightarrow B, A = \\{1, 2, 3\\}, x^{2}
"""

    theorem_proof_rule = """
[ANSWER KEY COMPLETENESS RULE - CRITICAL]
- "Refer Textbook" ([பாடநூல் பார்க்கவும்]) is allowed ONLY for:
  * Theorem PROOFS (தேற்றத்தை நிரூபிக்கவும்)
  * Definitions (வரையறு)
  * Compass/ruler CONSTRUCTION steps (வரைக/அமைக்க)
- For ALL other questions - numerical problems, equations, statistics,
  probability, coordinate geometry, mensuration, graph calculations -
  the Answer Key MUST contain COMPLETE STEP-BY-STEP working with the final answer.
  This is especially MANDATORY for compulsory questions and Part III/IV
  high-mark questions. A numerical question answered with Refer-Textbook is INVALID.
- For graph questions: Answer Key must include the table of values (x, y points)
  and the final answer read from the graph.

[ARITHMETIC ACCURACY RULE - CRITICAL]
- Before writing any Answer Key calculation, VERIFY every arithmetic step.
- For STATISTICS questions (mean, standard deviation, variance):
  * Choose data values so that the sum, sum of squares, and mean are CLEAN numbers.
  * Show working explicitly: n = ..., sum = ..., mean = ..., variance = ..., SD = ...
    with each intermediate value actually computed (not just formula with symbols).
  * The final SD should be a clean value (whole number or one decimal).
- If you cannot verify a calculation is correct, redesign the question with simpler numbers.

[OUTPUT HYGIENE RULE - ABSOLUTELY CRITICAL]
- This output is a FINAL PRINTED EXAM DOCUMENT given directly to students and teachers.
- NEVER include any meta-commentary, thinking process, or self-correction text such as:
  "Let me try a different approach", "Re-checking", "My initial analysis was correct",
  "This is complex", "Re-design the question", "நான் ... கணக்கிடுகிறேன்",
  "குறிப்பு: ... தெளிவாக குறிப்பிடப்படவில்லை" or ANY similar draft/verification notes.
- Output ONLY the polished final content. Any visible reasoning = INVALID output.

[QUESTION-ANSWER CONSISTENCY RULE - ABSOLUTELY CRITICAL]
- The question printed in the paper and the question solved in the Answer Key
  MUST be 100% IDENTICAL. NEVER solve a modified version in the Answer Key.
- DESIGN QUESTIONS BACKWARDS: first choose the clean final answer, then build
  the question from it. Examples:
  * Perfect-square polynomial question: FIRST pick a quadratic like (2x²-3x+2),
    square it, and use THAT expansion as the question polynomial. This guarantees
    the question is solvable.
  * Square root / factorisation: construct from known factors.
  * Statistics: pick data whose mean and SD are clean by design.
- If while writing the Answer Key you discover the question is unsolvable or messy,
  you MUST go back and replace the QUESTION itself with the corrected version,
  so the printed question matches its solution perfectly.

[STATISTICS TWO-DATASET RULE]
- If a statistics question gives two datasets (e.g. income x and expenditure y),
  the question must state EXPLICITLY what to compute:
  e.g. "இரு தரவுத்தொகுப்புகளின் மாறுபாட்டுக் கெழுக்களை (C.V.) கணக்கிட்டு ஒப்பிடுக" -
  and the Answer Key must then compute BOTH datasets fully.
- Never give two datasets and compute only one.

[MENSURATION VALUES RULE]
- When π = 22/7 is used, choose radius/height values that are multiples of 7
  so answers come out as whole numbers (e.g. r=7, r=14, r=21, h=7).
- Express intermediate fractions exactly (33000/7) only if they cancel later;
  final answers must be clean numbers.
"""

    quality_rules = """
[QUESTION PAPER QUALITY RULES - MANDATORY]
1. THEOREM DIAGRAMS ARE COMPULSORY:
   - EVERY theorem proof question (தேற்றம் எழுதி நிரூபிக்கவும்) MUST have its diagram tag.
   - தேல்ஸ் தேற்றம் → [DRAW_THALES] is NON-NEGOTIABLE. Never skip it.
   - A theorem question WITHOUT its diagram tag is an INVALID output.

2. CONSTRUCTION QUESTIONS - STRICT SCOPE:
   - The note "(உரிய அளவுகளுடன் வரைபடம் வரைந்து காட்டவும்)" must be added ONLY to
     COMPASS-AND-RULER CONSTRUCTION questions - i.e. questions whose main verb is
     "வரைக" or "அமைக்க" (construct/draw).
   - NEVER add this note to: prove questions (நிரூபி), calculation questions (காண்க),
     theorem questions, graph questions, or any non-construction question.
   - EVERY construction question MUST state ALL required measurements explicitly
     in the question text (side lengths, radius, scale factor etc). A construction
     question without complete measurements is INVALID - redesign it with numbers.

3. QUESTION CLARITY:
   - NEVER write vague questions like "வடிவொத்த முக்கோணங்கள் வரைய பயன்படும் தேற்றத்தை எழுதுக".
   - Always name the EXACT theorem: e.g. "அடிப்படை விகிதசம தேற்றத்தை (தேல்ஸ் தேற்றம்) எழுதி நிரூபிக்கவும்".
   - If a question asks about angle bisector, specify internal (உட்புற) or external (வெளிப்புற)
     AND state clearly where the intersection point lies (e.g. "D என்பது BC-ன் நீட்டிப்பில் உள்ளது").

4. MULTI-PART QUESTIONS:
   - If a question has two tasks (e.g. "நிரூபிக்கவும்" + "மதிப்பு காண்க"),
     split into (அ) and (ஆ) sub-parts explicitly.

5. ANSWER KEY CONSISTENCY:
   - The ratio/expression in the Answer Key MUST be IDENTICAL to the one asked in the question.
   - If question asks AO/BO = CO/DO, answer key must use AO/BO = CO/DO. Never switch to a different
     equivalent form. Never write "(அல்லது ...)" alternate forms.

6. REALISTIC VALUES:
   - Design numerical questions so answers are clean values (whole numbers or simple decimals
     like 4.5). Avoid awkward answers like 0.4 or 9.6 in geometry length problems.

7. MCQ QUALITY STANDARD:
   - Avoid trivial rote-memory MCQs (e.g. asking just the count of elements in a 3x2 matrix).
   - Each MCQ should require at least one step of thinking or calculation.
   - Mix: 40% concept application, 40% one-step calculation, maximum 20% pure recall.

8. BILINGUAL TECHNICAL TERMS:
   - For key technical concepts in Tamil medium papers, add the English term in
     brackets on first use: "முடிவிலா எண்ணிக்கையிலான தீர்வுகள் (infinitely many solutions)",
     "திட்ட விலக்கம் (Standard Deviation)". This matches TN Board pattern.
"""

    pyq_tagging_rule = """
[PREVIOUS YEAR QUESTION (PYQ) TAGGING RULE - TN BOARD PUBLIC EXAMS]
- Scope: Apply ONLY to 2-Mark, 5-Mark, and 8-Mark/Long-answer questions.
- Do NOT apply to 1-Mark MCQ questions under any circumstance.
- If a generated question appeared in a Tamil Nadu State Board PUBLIC EXAM
  in the last 4 years (2022, 2023, 2024, 2025), append the exam session tag
  at the END of that question in brackets.
- Tag format (Tamil papers):  (செப் 2024) / (ஏப்ரல் 2023) / (ஜூன் 2022)
- Tag format (English papers): (Sep 2024) / (Apr 2023) / (Jun 2022)
- Multiple appearances: list all, comma-separated: (ஏப்ரல் 2023, செப் 2024)
- CRITICAL ACCURACY RULE: Tag ONLY if you are HIGHLY CONFIDENT the question
  appeared in that specific exam session. If uncertain about the year or month,
  DO NOT tag at all. An untagged repeated question is acceptable;
  a WRONGLY tagged question is NOT acceptable.
- Do NOT invent or guess exam sessions. Do NOT tag more than 40% of questions.
"""

    no_latex_rule = """
[ABSOLUTE NO-LATEX FORMATTING RULE]
- NEVER use LaTeX syntax in the output. This is a Word document, not LaTeX.
- BANNED: \\{ \\} \\rightarrow \\leftarrow \\times \\div \\le \\ge \\ne \\in \\subset \\cup \\cap \\sqrt{} \\frac{}{} \\pi \\theta \\alpha \\beta \\angle \\triangle \\cdot $...$ etc.
- USE plain Unicode symbols instead:
  * Sets: {1, 2, 3} — plain curly braces WITHOUT backslash
  * Arrow: → (f: A → B)
  * Multiply: × | Divide: ÷ | Not equal: ≠ | Less/greater equal: ≤ ≥
  * Element of: ∈ | Subset: ⊂ | Union: ∪ | Intersection: ∩
  * Square root: √2, √(x+1) | Fraction: a/b or (x+1)/(x-2)
  * Pi: π | Theta: θ | Degree: 45° | Angle: ∠ABC | Triangle: △ABC
  * Squared: x² | Cubed: x³ | Subscript: a₁, a₂, aₙ
- Example CORRECT: "A = {1, 2, 3, 4}, B = {2, 5, 8} மற்றும் f: A → B ஆனது f(x) = 3x - 1"
- Example WRONG: "A = \\{1, 2, 3, 4\\}, f: A \\rightarrow B"
"""
    lang_override = ""
    if force_english:
        lang_override = ("\n[TOP PRIORITY LANGUAGE OVERRIDE] The ENTIRE question paper — every question, "
                         "instruction, part heading and option — MUST be written in PURE ENGLISH only. "
                         "Do NOT use any Tamil words. Mathematical/technical terms stay in English.")
    elif force_tamil:
        lang_override = ("\n[TOP PRIORITY LANGUAGE OVERRIDE] The ENTIRE question paper — every question, "
                         "instruction, part heading and option — MUST be written in PURE TAMIL only "
                         "(கணித சின்னங்கள்/எண்கள் தவிர). Do NOT write questions in English.")

    return f"Subject: {subject}\nLessons: {lessons_str}\nExam Type: {exam_type}\nTotal Marks: {total_marks}\nTime: {exam_time}\nMode: {exam_mode}\n{difficulty_directive}\n{blueprint_desc}\n{header_format}\n{option_format}\n{lang_instruction}{lang_override}\n{subject_blueprint_rules}\n{no_latex_rule}\n{quality_rules}\n{theorem_proof_rule}\n{pyq_tagging_rule}\n{no_latex_rule}\n=== ANSWER KEY ==="


# ==========================================
# LaTeX → Normal Text Converter
# ==========================================
LATEX_REPLACEMENTS = [
    (r'\\rightarrow', '→'), (r'\\to\b', '→'), (r'\\leftarrow', '←'),
    (r'\\Rightarrow', '⇒'), (r'\\leftrightarrow', '↔'),
    (r'\\in\b', '∈'), (r'\\notin', '∉'), (r'\\subseteq', '⊆'), (r'\\subset', '⊂'),
    (r'\\cup', '∪'), (r'\\cap', '∩'), (r'\\emptyset', '∅'), (r'\\phi', 'φ'),
    (r'\\times', '×'), (r'\\div', '÷'), (r'\\pm', '±'),
    (r'\\cdot', '·'), (r'\\neq', '≠'), (r'\\ne\b', '≠'),
    (r'\\leq', '≤'), (r'\\le\b', '≤'), (r'\\geq', '≥'), (r'\\ge\b', '≥'),
    (r'\\approx', '≈'), (r'\\equiv', '≡'), (r'\\infty', '∞'),
    (r'\\angle', '∠'), (r'\\triangle', '△'),
    (r'\\perp', '⊥'), (r'\\parallel', '∥'), (r'\\cong', '≅'), (r'\\sim\b', '∼'),
    (r'\\degree', '°'), (r'\\circ\b', '°'),
    (r'\\alpha', 'α'), (r'\\beta', 'β'), (r'\\gamma', 'γ'), (r'\\delta', 'δ'),
    (r'\\theta', 'θ'), (r'\\lambda', 'λ'), (r'\\mu\b', 'μ'), (r'\\pi\b', 'π'),
    (r'\\sigma', 'σ'), (r'\\omega', 'ω'), (r'\\Delta', 'Δ'), (r'\\Sigma', 'Σ'),
    (r'\\mathbb\{R\}', 'ℝ'), (r'\\mathbb\{N\}', 'ℕ'), (r'\\mathbb\{Z\}', 'ℤ'), (r'\\mathbb\{Q\}', 'ℚ'),
]

SUPERSCRIPT_MAP = {'0':'⁰','1':'¹','2':'²','3':'³','4':'⁴','5':'⁵','6':'⁶','7':'⁷','8':'⁸','9':'⁹','n':'ⁿ','-':'⁻','+':'⁺'}
SUBSCRIPT_MAP   = {'0':'₀','1':'₁','2':'₂','3':'₃','4':'₄','5':'₅','6':'₆','7':'₇','8':'₈','9':'₉','n':'ₙ','-':'₋','+':'₊'}

def latex_to_normal(text):
    """LaTeX symbols → normal Unicode text"""
    text = re.sub(r'\\d?frac\{([^{}]+)\}\{([^{}]+)\}', r'(\1/\2)', text)
    text = re.sub(r'\\sqrt\{([^{}]+)\}', r'√(\1)', text)
    text = re.sub(r'\\sqrt\s*', '√', text)
    for pattern, repl in LATEX_REPLACEMENTS:
        text = re.sub(pattern, repl, text)
    text = re.sub(r'\^\{([^{}]+)\}', lambda m: ''.join(SUPERSCRIPT_MAP.get(c, c) for c in m.group(1)), text)
    text = re.sub(r'\^(\d)', lambda m: SUPERSCRIPT_MAP.get(m.group(1), m.group(1)), text)
    text = re.sub(r'_\{([^{}]+)\}', lambda m: ''.join(SUBSCRIPT_MAP.get(c, c) for c in m.group(1)), text)
    text = re.sub(r'_(\d)', lambda m: SUBSCRIPT_MAP.get(m.group(1), m.group(1)), text)
    text = text.replace('\\{', '{').replace('\\}', '}')
    text = text.replace('$', '')
    text = re.sub(r'\\(?=[a-zA-Z])', '', text)  # strip leftover backslashes
    return text

def set_cell_margins(cell, **kwargs):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m in ["top", "start", "bottom", "end"]:
        if m in kwargs:
            node = OxmlElement(f"w:{m}")
            node.set(qn('w:w'), str(kwargs[m]))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
    tcPr.append(tcMar)

def add_solid_line(doc):
    p = doc.add_paragraph()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    p.paragraph_format.element.get_or_add_pPr().append(pBdr)

def write_markdown_to_word(doc, text):
    text = latex_to_normal(text)
    for line in text.split('\n'):
        line = line.strip()
        if not line:
            continue
        draw_match = re.search(r'\[DRAW_(THALES|BPT|EXT_BISECTOR|EXTERNAL_BISECTOR|ANGLE_BISECTOR|BISECTOR|TWO_TANGENTS|TWO_TANGENT|PYTHAGORAS|RIGHT_TRIANGLE|TANGENT|TRIANGLE|SQUARE|RECTANGLE|CIRCLE|SEMICIRCLE)[:\s]*(.*?)\]', line, re.IGNORECASE)
        if draw_match:
            shape_type = draw_match.group(1)
            label_text = draw_match.group(2)
            try:
                img_buf = generate_geometry_image(shape_type, label_text)
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.add_run().add_picture(img_buf, width=Inches(2.5))
            except Exception as e:
                doc.add_paragraph(f"[Error loading diagram: {e}]")
            continue
        clean_line = line.replace('*', '').replace('$', '').strip()
        if "பகுதி" in clean_line or "PART" in clean_line.upper():
            marks_match = re.search(r'\(?\d+\s*[xX*]\s*\d+\s*=\s*\d+\)?', clean_line)
            if marks_match:
                calc_str  = marks_match.group(0)
                title_str = clean_line.replace(calc_str, "").strip(":- ")
                table = doc.add_table(rows=1, cols=2)
                c1, c2 = table.rows[0].cells
                c1.paragraphs[0].add_run(title_str).bold = True
                c2.paragraphs[0].add_run(f"({calc_str})" if not calc_str.startswith("(") else calc_str).bold = True
                c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                set_cell_margins(c1, top=0, bottom=0, start=0, end=0)
                set_cell_margins(c2, top=0, bottom=0, start=0, end=0)
                continue
        option_markers = ["அ)", "ஆ)", "இ)", "ஈ)", "a)", "b)", "c)", "d)"]
        if any(marker in clean_line for marker in option_markers):
            raw_parts = re.split(r'(அ\)|ஆ\)|இ\)|ஈ\)|a\)|b\)|c\)|d\))', clean_line)
            parts = []
            current = ""
            for chunk in raw_parts:
                if chunk in option_markers:
                    if current.strip():
                        parts.append(current.strip())
                    current = chunk + " "
                else:
                    current += chunk
            if current.strip():
                parts.append(current.strip())
            if parts:
                table = doc.add_table(rows=1, cols=len(parts))
                for idx, opt in enumerate(parts):
                    cell = table.cell(0, idx)
                    cell.paragraphs[0].add_run(opt.replace("*", ""))
                    set_cell_margins(cell, top=0, bottom=0, start=0, end=0)
                continue
        p = doc.add_paragraph()
        if re.match(r'^\d+\.', clean_line):
            p.paragraph_format.left_indent       = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.25)
        parts = re.split(r'\*\*(.*?)\*\*', line)
        for i, part in enumerate(parts):
            run = p.add_run(part.replace('$', ''))
            if i % 2 == 1:
                run.bold = True

def create_professional_docx(ai_response, school_name, class_val, subject_val, exam_type, time_val, marks_val):
    doc = Document()
    section = doc.sections[0]
    section.page_width = section.page_height = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = section.right_margin = section.top_margin = section.bottom_margin = Inches(0.5)
    style = doc.styles['Normal']
    style.font.name = 'Nirmala UI'
    h_school = doc.add_paragraph(style='Normal')
    h_school.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h_school.add_run(school_name.upper()).bold = True
    h_school.runs[0].font.size = Pt(15)
    table = doc.add_table(rows=2, cols=2)
    def format_cell(cell, text, align_right=False):
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if align_right else WD_ALIGN_PARAGRAPH.LEFT
        p.add_run(text).bold = True
        set_cell_margins(cell, top=0, bottom=0, start=0, end=0)
    format_cell(table.cell(0, 0), f"Class : {class_val}")
    format_cell(table.cell(0, 1), f"Time : {time_val}", align_right=True)
    format_cell(table.cell(1, 0), f"Subject : {subject_val}")
    format_cell(table.cell(1, 1), f"Marks : {marks_val}", align_right=True)
    add_solid_line(doc)
    parts = ai_response.split("=== ANSWER KEY ===")
    write_markdown_to_word(doc, parts[0].strip())
    return doc

# ==========================================
# COMPACT HEADER (profile + logout inline, no left sidebar)
# ==========================================
_plan_label = "👑 Premium" if user_plan != "free" else "👑 Free"
_first_name = (user_teacher or user_name or "Teacher").split()[0]

hdr_l, hdr_r = st.columns([3, 1])
with hdr_l:
    st.markdown(f"""
    <div class="pmp-header" style="padding:16px 22px;margin-bottom:14px;">
        <div>
            <h1 style="font-size:20px !important;">வரவேற்கிறோம், <span class="accent">{_first_name}</span>! 👋</h1>
            <p style="font-size:12.5px;">🏫 {user_school or ''} · 📞 {current_user.get('mobile','')}</p>
        </div>
        <div class="pmp-badge">{('⭐ Premium' if is_premium else (f'🎁 Trial · {trial_days_left} நாள்' if trial_days_left is not None else '🎁 Trial'))}</div>
    </div>
    """, unsafe_allow_html=True)
with hdr_r:
    st.markdown("<div style='height:14px;'></div>", unsafe_allow_html=True)
    if st.button("🚪 Logout", use_container_width=True):
        st.session_state.pop("logged_in_user", None); st.session_state.pop("intro_shown", None)
        st.rerun()
    with st.expander("✏️ Profile திருத்து", expanded=not (user_school and current_user.get("mobile"))):
        e_school  = st.text_input("பள்ளி", value=user_school or "", key="edit_school")
        e_teacher = st.text_input("ஆசிரியர்", value=user_teacher or "", key="edit_teacher")
        e_mobile  = st.text_input("மொபைல்", value=current_user.get("mobile", "") or "", max_chars=10, key="edit_mobile")
        if st.button("💾 சேமி", key="edit_save", use_container_width=True):
            if not e_school.strip():
                st.error("பள்ளி பெயரை நிரப்பவும்")
            elif e_mobile.strip() and not (e_mobile.strip().isdigit() and len(e_mobile.strip()) == 10):
                st.error("சரியான 10 இலக்க மொபைல்")
            else:
                ok = update_user_profile(user_id, e_school, e_teacher, e_mobile, email=user_email)
                if ok:
                    # Re-fetch from DB to confirm it actually persisted
                    _chk = fetch_user_by_email(user_email)
                    if _chk and (_chk.get("school_name") or "").strip() == e_school.strip():
                        st.session_state["logged_in_user"] = _chk
                        st.success("✅ சேமிக்கப்பட்டது!")
                        st.rerun()
                    else:
                        st.error("⚠️ DB-ல் சேமிக்க முடியலை. Neon-ல் school_name/teacher_name/mobile columns இருக்கா என்று சரிபார்க்கவும்.")
                else:
                    st.error("⚠️ சேமிப்பில் பிழை. DB connection சரிபார்க்கவும்.")

# ==========================================
# MAIN APP TABS
# ==========================================
tab1, tab2, tab3 = st.tabs([
    "🎓 வினாத்தாள் தயாரிப்பு",
    "📝 விடைத்தாள் திருத்தம்",
    "🛠️ கேள்வி வங்கி மேலாண்மை",
])

with tab1:
    df = load_data()
    if not df.empty:
        # Row 1: School, Class, Subject, Exam Type (4 columns)
        r1c1, r1c2, r1c3, r1c4 = st.columns(4)
        with r1c1:
            school_name = st.text_input("School Name", value=st.session_state["logged_in_user"].get("school_name") or user_school or "ABC SCHOOL")
        with r1c2:
            class_val = st.selectbox("Class", ["10"])
        with r1c3:
            subject_list = df['Subject'].unique()
            subject_val  = st.selectbox("Subject", subject_list)
        with r1c4:
            exam_type = st.selectbox("Exam Type", ["Unit Test", "Quarterly Exam", "Half-Yearly Exam", "Annual Exam"])

        # Row 2: Time, Marks, Difficulty, Mode (4 columns)
        r2c1, r2c2, r2c3, r2c4 = st.columns(4)
        with r2c1:
            time_val = st.selectbox("Time (நேரம்)", ["1.00 Hour", "2.00 Hours", "3.00 Hours"], index=2)
        with r2c2:
            marks_val = st.number_input("Total Marks", value=100, step=1)
        with r2c3:
            diff_level = st.selectbox("கடினத்தன்மை", ["எளிமை (Easy)", "நடுத்தரம் (Medium)", "கடினம் (Hard)"], index=1)
        with r2c4:
            exam_mode = st.selectbox("Exam Mode", ["🏛️ Public", "🏫 School Elite"])

        # Row 3: Language toggle (left) + Lessons multiselect (right)
        r3c1, r3c2 = st.columns([1, 2])
        with r3c1:
            paper_lang = st.radio("மொழி (Language)", ["தமிழ் (Tamil)", "English"], horizontal=True)
        with r3c2:
            lesson_list      = df[df['Subject'] == subject_val]['Lesson'].unique()
            selected_lessons = st.multiselect("பாடங்களைத் தேர்ந்தெடுக்கவும்", lesson_list)

        st.markdown("<hr style='margin:8px 0;'>", unsafe_allow_html=True)

        is_eng = "english" in subject_val.lower() or "ஆங்கிலம்" in subject_val.lower()
        is_soc = "social"  in subject_val.lower() or "சமூக"     in subject_val.lower()
        bp     = get_blueprint_defaults(marks_val, is_social=is_soc, is_english=is_eng)

        # ===== Part selector cards =====
        _pcard = st.container(border=True)
        with _pcard:
            st.markdown("""
            <div style='display:flex;align-items:center;gap:10px;margin-bottom:4px;'>
                <div style='width:34px;height:34px;border-radius:9px;background:#eef1fd;display:flex;
                            align-items:center;justify-content:center;font-size:17px;'>📋</div>
                <div><h3 style='margin:0;'>வினா வடிவமைப்பு பிரிவு</h3>
                <span style='color:#5a6782;font-size:12px;'>Paper-ல் சேர்க்க வேண்டிய வினா வகைகளைத் தேர்ந்தெடுக்கவும்</span></div>
            </div>
            """, unsafe_allow_html=True)

            pc1, pc2, pc3, pc4 = st.columns(4)
            with pc1:
                show_p1 = st.checkbox("பகுதி I · 1-Mark", value=True)
            with pc2:
                show_p2 = st.checkbox("பகுதி II · 2-Mark", value=True)
            with pc3:
                show_p3 = st.checkbox("பகுதி III · 5-Mark", value=True)
            with pc4:
                show_p4 = st.checkbox("பகுதி IV · நெடுவினா", value=True)

        # ===== Marks details + live donut =====
        _mcard = st.container(border=True)
        with _mcard:
            marks_col, donut_col = st.columns([2, 1])

        with marks_col:
            st.markdown("#### ⚙️ மதிப்பெண் விவரங்கள்")

            # Row 1: 1-mark (full width)
            p1_ask = st.number_input("1-மார்க் வினாக்கள்", min_value=0, max_value=30,
                                     value=int(bp["p1"]) if show_p1 else 0, step=1, disabled=not show_p1)

            # Row 2: 2-mark Given (left) + Answer (right)
            r2l, r2r = st.columns(2)
            with r2l:
                p2_get = st.number_input("2-மார்க் கொடுக்க (Given)", min_value=0, max_value=30,
                                         value=int(bp["p2g"]) if show_p2 else 0, step=1, disabled=not show_p2)
            with r2r:
                p2_ask = st.number_input("2-மார்க் எழுத (Answer)", min_value=0, max_value=30,
                                         value=int(bp["p2a"]) if show_p2 else 0, step=1, disabled=not show_p2)

            # Row 3: 5-mark Given (left) + Answer (right)
            r3l, r3r = st.columns(2)
            with r3l:
                p3_get = st.number_input("5-மார்க் கொடுக்க (Given)", min_value=0, max_value=30,
                                         value=int(bp["p3g"]) if show_p3 else 0, step=1, disabled=not show_p3)
            with r3r:
                p3_ask = st.number_input("5-மார்க் எழுத (Answer)", min_value=0, max_value=30,
                                         value=int(bp["p3a"]) if show_p3 else 0, step=1, disabled=not show_p3)

            # Row 4: நெடுவினா mark value + Given + Answer
            r4a, r4b, r4c = st.columns(3)
            with r4a:
                p4_val = st.selectbox("நெடுவினா மதிப்பெண்", [5, 8, 10],
                                      index=1 if is_eng or is_soc or marks_val == 100 else 0, disabled=not show_p4)
            with r4b:
                p4_get = st.number_input("நெடுவினா கொடுக்க", min_value=0, max_value=20,
                                         value=int(bp["p4g"]) if show_p4 else 0, step=1, disabled=not show_p4)
            with r4c:
                p4_ask = st.number_input("நெடுவினா எழுத", min_value=0, max_value=20,
                                         value=int(bp["p4a"]) if show_p4 else 0, step=1, disabled=not show_p4)

        total_calculated = (p1_ask * 1) + (p2_ask * 2) + (p3_ask * 5) + (p4_ask * p4_val)
        can_generate     = total_calculated == marks_val

        # ===== Compact summary (no donut, saves space) =====
        with donut_col:
            st.markdown("##### 📊 Blueprint Summary")
            _bal_color = "#1b9e5a" if can_generate else "#9a6b00"
            _bal_bg    = "#e8f5ec" if can_generate else "#fff6df"
            st.markdown(
                f"<div style='background:{_bal_bg};border-radius:12px;padding:10px 14px;text-align:center;margin-bottom:8px;'>"
                f"<span style='font-family:Sora,sans-serif;font-size:30px;font-weight:800;color:#0a1f44;'>{total_calculated}</span>"
                f"<span style='font-size:15px;color:#5a6782;'> / {marks_val}</span></div>",
                unsafe_allow_html=True)

            rows = [
                ("🟢", "1-மார்க்",        p1_ask,  "" ),
                ("🔵", "2-மார்க் Given",  p2_get,  "" ),
                ("🟣", "2-மார்க் Ans",    p2_ask,  "" ),
                ("🟠", "5-மார்க் Given",  p3_get,  "" ),
                ("🩷", "5-மார்க் Ans",    p3_ask,  "" ),
                ("🟦", "நெடுவினா Given",  p4_get,  "" ),
                ("🔴", "நெடுவினா Ans",    p4_ask,  "" ),
            ]
            html = "<div style='font-size:13px;line-height:1.5;'>"
            for dot, lbl, cnt, _ in rows:
                if cnt <= 0:
                    continue
                html += (f"<div style='display:flex;justify-content:space-between;padding:3px 0;border-bottom:1px solid #f0f2f8;'>"
                         f"<span style='color:#5a6782;'>{dot} {lbl}</span>"
                         f"<b style='color:#0f1a30;'>{cnt}</b></div>")
            html += "</div>"
            st.markdown(html, unsafe_allow_html=True)

            if can_generate:
                st.markdown("<div style='background:#e8f5ec;border:1px solid #22c55e;border-radius:10px;padding:8px;margin-top:8px;text-align:center;'>"
                            "<b style='color:#1b9e5a;'>✅ Balanced</b></div>", unsafe_allow_html=True)
            else:
                st.markdown(f"<div style='background:#fff6df;border:1px solid #c9a227;border-radius:10px;padding:8px;margin-top:8px;text-align:center;'>"
                            f"<b style='color:#9a6b00;'>⚠️ சமப்படுத்தவும்</b></div>", unsafe_allow_html=True)

        st.markdown("<hr style='margin:6px 0;'>", unsafe_allow_html=True)
        gen_mode = st.radio(
            "வினாத்தாள் உருவாக்கும் முறை",
            ["🤖 AI Auto-Generate (தற்போதைய முறை)", "📚 கேள்வி வங்கியில் இருந்து தேர்ந்தெடு (Book Back / Exercise / Example)"],
            horizontal=False,
        )
        bank_mode = gen_mode.startswith("📚")

        parts_meta = [
            {"key": "p1", "label": "பகுதி I",   "mark_type": "1M",   "mark": 1,     "given": int(p1_ask), "answer": int(p1_ask), "show": show_p1, "note": ""},
            {"key": "p2", "label": "பகுதி II",  "mark_type": "2M",   "mark": 2,     "given": int(p2_get), "answer": int(p2_ask), "show": show_p2, "note": (f"ஏதேனும் {int(p2_ask)} கேள்விகளுக்கு மட்டும் விடையளிக்கவும்" if p2_ask != p2_get else "")},
            {"key": "p3", "label": "பகுதி III", "mark_type": "5M",   "mark": 5,     "given": int(p3_get), "answer": int(p3_ask), "show": show_p3, "note": (f"ஏதேனும் {int(p3_ask)} கேள்விகளுக்கு மட்டும் விடையளிக்கவும்" if p3_ask != p3_get else "")},
            {"key": "p4", "label": "பகுதி IV",  "mark_type": "LONG", "mark": int(p4_val), "given": int(p4_get), "answer": int(p4_ask), "show": show_p4, "note": (f"ஏதேனும் {int(p4_ask)} கேள்விகளுக்கு மட்டும் விடையளிக்கவும்" if p4_ask != p4_get else "")},
        ]

        if bank_mode:
            st.markdown("### 📚 கேள்வி வங்கி — தேர்ந்தெடுக்கவும்")
            if not selected_lessons:
                st.warning("⚠️ முதலில் மேலே பாடங்களைத் தேர்ந்தெடுக்கவும்!")
            else:
                if st.button("🔄 கேள்வி வங்கியை ஏற்று (Load Question Bank)", use_container_width=True):
                    with st.spinner("⏳ கேள்வி வங்கியில் இருந்து ஏற்றப்படுகிறது..."):
                        pool = {}
                        for part in parts_meta:
                            if not part["show"] or part["given"] <= 0:
                                continue
                            merged = []
                            for lesson in selected_lessons:
                                merged.extend(get_or_build_bank(subject_val, lesson, part["mark_type"], min_count=part["given"]))
                            pool[part["key"]] = merged
                        st.session_state["bank_pool"] = pool
                        st.session_state["bank_pool_subject"] = subject_val
                        _total_loaded = sum(len(v) for v in pool.values())
                        if _total_loaded > 0:
                            st.success(f"✅ கேள்வி வங்கி தயார்! ({_total_loaded} கேள்விகள் கிடைத்தன)")
                        else:
                            st.warning("⚠️ தேர்ந்தெடுத்த பாடங்களுக்கு கேள்வி வங்கியில் கேள்விகள் இல்லை. "
                                       "'🛠️ கேள்வி வங்கி மேலாண்மை' Tab-ல் Excel Import செய்யப்பட்டதா என்று பாருங்க.")

                if "bank_pool" in st.session_state and st.session_state.get("bank_pool_subject") == subject_val:
                    for part in parts_meta:
                        if not part["show"] or part["given"] <= 0:
                            continue
                        pool = st.session_state["bank_pool"].get(part["key"], [])
                        if not pool:
                            continue
                        with st.expander(f'{part["label"]} — {MARK_TYPE_LABELS.get(part["mark_type"], part["mark_type"])} (தேவை: {part["given"]})', expanded=False):
                            col_left, col_right = st.columns([3, 2])
                            selected_items = []
                            with col_left:
                                st.caption("⬅️ கிடைக்கும் கேள்விகள் — tick செய்யவும்")
                                for it in pool:
                                    chk_key = f'chk_{part["key"]}_{it["id"]}'
                                    ref = it.get("reference", "") or ""
                                    tag = f'[{it.get("qtype","")}] '
                                    st.checkbox(f'{tag}{it["question_text"]}', key=chk_key)
                                    if ref:
                                        st.caption(f'📌 {ref}')
                            with col_right:
                                st.caption("➡️ தேர்ந்தெடுக்கப்பட்டவை")
                                for it in pool:
                                    chk_key = f'chk_{part["key"]}_{it["id"]}'
                                    if st.session_state.get(chk_key):
                                        selected_items.append(it)
                                        st.markdown(f'✅ {it["question_text"][:60]}{"..." if len(it["question_text"])>60 else ""}')
                                got = len(selected_items)
                                need = part["given"]
                                if got == need:
                                    st.success(f'{got}/{need} தேர்ந்தெடுக்கப்பட்டது ✅')
                                elif got > need:
                                    st.error(f'{got}/{need} — {got-need} அதிகமாக உள்ளது, சிலவற்றை நீக்கவும்')
                                else:
                                    st.warning(f'{got}/{need} தேர்ந்தெடுக்கப்பட்டது')

        _gb1, _gb2, _gb3 = st.columns([1, 2, 1])
        with _gb2:
            gen_clicked = st.button("🚀 Generate PRO Question Paper", use_container_width=True, type="primary")
        if gen_clicked:
            if bank_mode:
                if not selected_lessons:
                    st.warning("⚠️ பாடங்களைத் தேர்ந்தெடுக்கவும்!")
                elif "bank_pool" not in st.session_state:
                    st.warning("⚠️ முதலில் 'கேள்வி வங்கியை ஏற்று' பொத்தானை அழுத்தவும்!")
                else:
                    parts_cfg = []
                    all_ok = True
                    for part in parts_meta:
                        if not part["show"] or part["given"] <= 0:
                            continue
                        pool = st.session_state["bank_pool"].get(part["key"], [])
                        chosen = [it for it in pool if st.session_state.get(f'chk_{part["key"]}_{it["id"]}')]
                        if len(chosen) != part["given"]:
                            all_ok = False
                        parts_cfg.append({
                            "label": part["label"], "mark": part["mark"],
                            "given": part["given"], "answer": part["answer"],
                            "note": part["note"], "items": chosen,
                        })
                    if not all_ok:
                        st.error("⚠️ ஒவ்வொரு பகுதியிலும் தேவையான எண்ணிக்கை கேள்விகளை சரியாக தேர்ந்தெடுக்கவும் (மேலே உள்ள எண்களைப் பார்க்கவும்).")
                    else:
                        with st.spinner("⏳ தேர்ந்தெடுக்கப்பட்ட கேள்விகளுடன் வினாத்தாள் தயாராகிறது..."):
                            assembled_text = assemble_paper_from_bank(parts_cfg)
                            increment_usage(user_id)
                            doc    = create_professional_docx(assembled_text, school_name, class_val, subject_val, exam_type, time_val, marks_val)
                            doc_io = io.BytesIO()
                            doc.save(doc_io)
                            st.session_state['docx_bytes'] = doc_io.getvalue()
                            st.success("✅ வினாத்தாள் தயாராகிவிட்டது! (கேள்வி வங்கியில் இருந்து)")
            elif can_generate and selected_lessons:
                with st.spinner("⏳ வினாத்தாள் தயாராகிறது..."):
                    blueprint_desc = f"- Part I: {p1_ask} Qs. - Part II: Given {p2_get}, Answer {p2_ask}. - Part III: Given {p3_get}, Answer {p3_ask}. - Part IV: Given {p4_get}, Answer {p4_ask}."
                    prompt = generate_prompt_v18(subject_val, selected_lessons, exam_type, time_val, marks_val, exam_mode, blueprint_desc, p1_ask, p2_ask, p3_ask, diff_level, paper_lang)
                    response = None
                    for attempt in range(4):
                        try:
                            response = client.models.generate_content(model='gemini-2.5-flash', contents=prompt)
                            break
                        except Exception as api_err:
                            if ("429" in str(api_err) or "503" in str(api_err)) and attempt < 3:
                                time.sleep((attempt + 1) * 4)
                            else:
                                st.error(f"சர்வர் பிழை: {api_err}")
                    if response:
                        increment_usage(user_id)
                        doc    = create_professional_docx(response.text, school_name, class_val, subject_val, exam_type, time_val, marks_val)
                        doc_io = io.BytesIO()
                        doc.save(doc_io)
                        st.session_state['docx_bytes'] = doc_io.getvalue()
                        st.success("✅ வினாத்தாள் தயாராகிவிட்டது!")
            elif not selected_lessons:
                st.warning("⚠️ பாடங்களைத் தேர்ந்தெடுக்கவும்!")

        if 'docx_bytes' in st.session_state:
            st.download_button(
                label="📥 Download Word File (.docx)",
                data=st.session_state['docx_bytes'],
                file_name=f"PMP_{subject_val}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
    else:
        st.warning("⚠️ lesson_master_v1_5.csv கோப்பு கிடைக்கவில்லை.")

with tab2:
    st.title("📝 AI Answer Sheet Evaluator")
    uploaded_file = st.file_uploader("விடைத்தாளைத் தேர்ந்தெடுக்கவும் (Image / PDF)", type=["png", "jpg", "jpeg", "pdf"])
    if uploaded_file is not None:
        file_type   = uploaded_file.name.split(".")[-1].lower()
        eval_payload= []
        if file_type == "pdf":
            st.info("📊 PDF கோப்பு கண்டறியப்பட்டது.")
            file_data = uploaded_file.read()
            pdf_part  = types.Part.from_bytes(data=file_data, mime_type="application/pdf")
            eval_payload.append(pdf_part)
        else:
            image = Image.open(uploaded_file)
            st.image(image, caption="Uploaded Handwriting Page", width=420)
            eval_payload.append(image)

        if st.button("🚀 Start AI Evaluation", use_container_width=True):
            with st.spinner("⏳ AI திருத்தி வருகிறது..."):
                eval_prompt = "You are an official TN Board Math Evaluator. Read handwriting or PDF pages and correct step-by-step. Write fractions as $\\frac{a}{b}$ inside single dollar signs. Respond in Tamil."
                eval_payload.append(eval_prompt)
                response = None
                for attempt in range(3):
                    try:
                        response = client.models.generate_content(model='gemini-2.5-flash', contents=eval_payload)
                        break
                    except Exception as eval_err:
                        if "429" in str(eval_err) and attempt < 2:
                            time.sleep(4)
                        else:
                            st.error(f"மதிப்பீட்டு சர்வர் பிழை: {eval_err}")
                if response:
                    st.markdown(response.text)

with tab3:
    st.title("🛠️ கேள்வி வங்கி மேலாண்மை")
    st.caption("இங்கு அனைத்து பாடங்களின் கேள்வி வங்கியையும் திருத்தலாம், நீக்கலாம், அல்லது Excel-ல் இருந்து bulk import பண்ணலாம்.")

    st.markdown("### 📥 Excel Template Download / Bulk Import")
    imp_col1, imp_col2 = st.columns(2)
    with imp_col1:
        try:
            with open("question_bank_template.xlsx", "rb") as f:
                st.download_button(
                    "📥 Excel Template Download பண்ணு",
                    data=f.read(),
                    file_name="question_bank_template.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True,
                )
        except FileNotFoundError:
            st.info("Template file app folder-ல் இல்லை — Claude கொடுத்த question_bank_template.xlsx-ஐ app repo-வில் சேர்க்கவும்.")
    with imp_col2:
        import_file = st.file_uploader("📤 நிரப்பிய Excel-ஐ Upload பண்ணு", type=["xlsx"], key="bank_import_upload")
        if import_file is not None:
            if st.button("🚀 DB-ல் Import பண்ணு", use_container_width=True):
                with st.spinner("⏳ Import ஆகிறது..."):
                    try:
                        xls = pd.ExcelFile(import_file)
                        # Prefer "Question_Bank" sheet; else use the first sheet
                        sheet = "Question_Bank" if "Question_Bank" in xls.sheet_names else xls.sheet_names[0]
                        df_import = pd.read_excel(xls, sheet_name=sheet)
                        # Validate expected columns
                        need = {"Subject", "Lesson", "Mark_Type", "QType", "Question", "Answer"}
                        if not need.issubset(set(df_import.columns)):
                            st.error(
                                f"⚠️ இந்த Excel-ல் தேவையான columns இல்லை. "
                                f"Claude கொடுத்த converted file (chapterXX_TNpattern.xlsx) upload செய்யவும் — "
                                f"original PMP EduAI file (Questions/Answers/Metadata sheets) அல்ல.\n\n"
                                f"கிடைத்த columns: {list(df_import.columns)[:8]}"
                            )
                        else:
                            inserted, skipped, err = bulk_import_bank_from_df(df_import)
                            if err:
                                st.error(f"பிழை: {err}")
                            else:
                                st.success(f"✅ {inserted} கேள்விகள் புதிதாக சேர்க்கப்பட்டது. {skipped} ஏற்கனவே இருந்ததால் தவிர்க்கப்பட்டது.")
                    except Exception as e:
                        st.error(f"⚠️ Import பிழை: {e}")

    st.markdown("---")
    st.markdown("### 🤖 Answer இல்லாத கேள்விகளுக்கு AI Solution நிரப்பு")
    st.caption("Exercise/Unit Exercise-ல் proof-type கேள்விகளுக்கு பொதுவா Answer காலியா இருக்கும் — இங்கு Gemini வெச்சு batch-ஆ நிரப்பலாம்.")
    af1, af2, af3 = st.columns(3)
    with af1:
        af_subjects = list_bank_subjects()
        af_subject = st.selectbox("பாடம் (filter)", ["அனைத்தும்"] + af_subjects, key="af_subject") if af_subjects else None
    with af2:
        af_lessons = list_bank_lessons(af_subject) if (af_subject and af_subject != "அனைத்தும்") else []
        af_lesson = st.selectbox("பாடப்பிரிவு (filter)", ["அனைத்தும்"] + af_lessons, key="af_lesson")
    with af3:
        af_limit = st.number_input("எத்தனை கேள்விகள் (max)", min_value=10, max_value=500, value=50, step=10, key="af_limit")

    if st.button("🤖 Missing Answers-ஐ AI வெச்சு நிரப்பு", use_container_width=True, type="primary"):
        subj_param = None if af_subject in (None, "அனைத்தும்") else af_subject
        lesson_param = None if af_lesson == "அனைத்தும்" else af_lesson
        missing = fetch_bank_missing_answers(subject=subj_param, lesson=lesson_param, limit=int(af_limit))
        if not missing:
            st.info("Answer இல்லாத கேள்விகள் எதுவும் கிடைக்கவில்லை (இந்த filter-ல்).")
        else:
            with st.spinner(f"⏳ {len(missing)} கேள்விகளுக்கு AI Solution generate ஆகிறது... (batches-ஆ நடக்கும், நேரம் ஆகலாம்)"):
                answers = generate_answers_batch_ai(missing, batch_size=10)
                filled = 0
                for item in missing:
                    ans = answers.get(item["id"])
                    if ans:
                        update_bank_question(item["id"], item["question_text"], ans, item["qtype"])
                        filled += 1
                st.success(f"✅ {filled}/{len(missing)} கேள்விகளுக்கு Answer நிரப்பப்பட்டது! (மேலே 'கேள்விகளை திருத்து' section-ல் refresh செய்து பாருங்க)")

    st.markdown("---")
    st.markdown("### ✏️ கேள்விகளை திருத்து / நீக்கு")
    all_subjects = list_bank_subjects()
    if not all_subjects:
        st.info("இன்னும் கேள்வி வங்கியில் எதுவும் இல்லை. Tab 1-ல் 'Load Question Bank' பண்ணி உருவாக்கவும், அல்லது மேலே Excel Import பண்ணவும்.")
    else:
        f1, f2, f3 = st.columns(3)
        with f1:
            filt_subject = st.selectbox("பாடம்", all_subjects, key="mgmt_subject")
        with f2:
            lessons_for_subject = list_bank_lessons(filt_subject) if filt_subject else []
            filt_lesson = st.selectbox("பாடப்பிரிவு", ["அனைத்தும்"] + lessons_for_subject, key="mgmt_lesson")
        with f3:
            filt_mark = st.selectbox("மார்க் வகை", ["அனைத்தும்", "1M", "2M", "5M", "LONG"], key="mgmt_mark")

        rows = fetch_bank_filtered(
            subject=filt_subject,
            lesson=None if filt_lesson == "அனைத்தும்" else filt_lesson,
            mark_type=None if filt_mark == "அனைத்தும்" else filt_mark,
        )
        st.caption(f"மொத்தம் {len(rows)} கேள்விகள் கிடைத்தது")

        for r in rows:
            with st.expander(f'[{r["mark_type"]}] {r["lesson"]} — {r["question_text"][:70]}{"..." if len(r["question_text"])>70 else ""}'):
                new_qtype = st.selectbox("QType", ["பின்புற வினா", "பயிற்சி", "எடுத்துக்காட்டு"],
                                          index=["பின்புற வினா", "பயிற்சி", "எடுத்துக்காட்டு"].index(r["qtype"]) if r["qtype"] in ["பின்புற வினா", "பயிற்சி", "எடுத்துக்காட்டு"] else 1,
                                          key=f'qtype_{r["id"]}')
                new_q = st.text_area("கேள்வி", value=r["question_text"], key=f'q_{r["id"]}', height=90)
                new_a = st.text_area("விடை", value=r["answer_text"], key=f'a_{r["id"]}', height=90)
                bc1, bc2 = st.columns(2)
                with bc1:
                    if st.button("💾 Save", key=f'save_{r["id"]}', use_container_width=True):
                        if update_bank_question(r["id"], new_q, new_a, new_qtype):
                            st.success("✅ Save ஆனது!")
                with bc2:
                    if st.button("🗑️ Delete", key=f'del_{r["id"]}', use_container_width=True):
                        if delete_bank_question(r["id"]):
                            st.success("🗑️ நீக்கப்பட்டது! Page-ஐ refresh பண்ணவும்.")
                            st.rerun()
